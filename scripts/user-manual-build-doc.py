from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path("Users Manaul Qacc")
DATA = ROOT / "data" / "manual-capture-results.json"
VIDEO_DATA = ROOT / "data" / "manual-video-results.json"
DOCX_OUT = ROOT / "ERPTH-Qacc-User-Manual-TH.docx"
PDF_OUT = ROOT / "ERPTH-Qacc-User-Manual-TH.pdf"
QA_OUT = ROOT / "QA-REPORT.md"


CHAPTER_GUIDANCE = {
    "เริ่มต้นใช้งาน": [
        "ใช้หน้านี้เพื่อตรวจสอบภาพรวมธุรกิจ เมนูหลัก และสถานะการเชื่อมต่อก่อนเริ่มทำงานประจำวัน",
        "ผู้ใช้สามารถค้นหาเมนูได้จากช่องค้นหาด่วน และใช้เมนูเพิ่มเติมเพื่อเข้าถึงงานรอง",
    ],
    "ข้อมูลหลัก": [
        "จัดการข้อมูลลูกค้า ผู้ขาย และสินค้า/บริการก่อนสร้างเอกสารขายหรือซื้อ",
        "ตรวจสอบสถานะ Active/Archive และข้อมูลภาษีให้ถูกต้อง เพราะข้อมูลส่วนนี้ถูกดึงไปใช้ในเอกสารธุรกรรม",
    ],
    "การขาย": [
        "เริ่มจากใบเสนอราคา/SO แล้วติดตามต่อเป็นใบแจ้งหนี้ ใบเสร็จ และเอกสารที่เกี่ยวข้อง",
        "หน้ารายละเอียดใช้ตรวจสถานะเอกสาร ยอดเงิน รายการสินค้า และปุ่มดำเนินการถัดไป",
    ],
    "การซื้อ": [
        "ใช้คำขอซื้อและใบสั่งซื้อเพื่อควบคุมกระบวนการจัดซื้อก่อนรับสินค้าและตรวจ vendor bill",
        "ถ้าหน้าใดไม่มีข้อมูลหรือแสดงข้อผิดพลาด ให้ตรวจสอบสิทธิ์งานจัดซื้อและ endpoint ของ backend ที่เชื่อมต่อ",
    ],
    "รายจ่ายและเอกสารตรวจสอบ": [
        "ใช้รายจ่ายสำหรับบันทึกรายการค่าใช้จ่าย และใช้กล่องตรวจสอบเอกสารสำหรับงานเอกสารที่เข้ามาจากช่องทางอัตโนมัติ",
        "Copilot/AI เป็นตัวช่วยตรวจ ไม่ควรถือเป็นการอนุมัติหรือบันทึกรายการแทนผู้ใช้",
    ],
    "บัญชี/ภาษี/e-Tax": [
        "ใช้รายงานบัญชีเพื่อดูภาพรวมและเจาะดูรายการทางบัญชีตามช่วงวันที่",
        "หน้า VAT/WHT/e-Tax และหน้าตั้งค่าผู้ดูแลระบบ ควรให้ผู้ใช้ที่รับผิดชอบบัญชีหรือผู้ดูแลระบบใช้งาน",
    ],
    "เครื่องมือ": [
        "เครื่องมือเหล่านี้ช่วยนำเข้า Excel, ใช้ Agent, ตั้งค่าการเชื่อมต่อ และออกแบบเอกสารใน Reports Studio",
        "บางฟังก์ชันต้องมี token หรือสิทธิ์เพิ่มเติมจาก backend จึงอาจแสดงข้อจำกัดตามสิทธิ์จริง",
    ],
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18 if level == 1 else 13)
    run.font.color.rgb = RGBColor(8, 61, 71) if level == 1 else RGBColor(28, 91, 103)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("- ")
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        t = p.add_run(item)
        t.font.name = "Arial"
        t.font.size = Pt(10.5)


def add_image(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        add_body(doc, f"[ไม่พบภาพ: {image_path.name}] {caption}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.4))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 108, 118)


def status_th(status: str, error: str | None = None) -> str:
    if error:
        return "มีข้อผิดพลาด"
    return {
        "loaded": "เปิดได้",
        "login": "หน้า Login",
        "empty-state": "เปิดได้ แต่ยังไม่มีข้อมูล",
        "error": "แสดง error จริง",
        "restricted": "จำกัดสิทธิ์",
        "not-found": "ไม่พบหน้า",
    }.get(status or "", status or "-")


def clean_text(text: str) -> str:
    replacements = {
        "Print Preview": "ตัวอย่างก่อนพิมพ์",
        "Drilldown": "ดูรายละเอียดเจาะลึก",
        "Partner Ledger": "รายงานลูกหนี้/เจ้าหนี้รายคู่ค้า",
        "Move Line": "รายการบัญชี",
        "Vendor Bill": "บิลผู้ขาย",
        "Delivery": "การส่งสินค้า",
        "Sale Order": "คำสั่งขาย",
        "General Ledger": "บัญชีแยกประเภท",
        "Aged Receivables": "อายุลูกหนี้",
        "Aged Payables": "อายุเจ้าหนี้",
        "Cash Book": "สมุดเงินสด",
        "Bank Book": "สมุดเงินฝากธนาคาร",
        "e-Tax Settings": "ตั้งค่า e-Tax",
        "Reports Studio": "สตูดิโอรายงาน",
        "Agent Dashboard": "แดชบอร์ด Agent",
        "Agent OCR": "Agent อ่านเอกสาร OCR",
        "Agent Expense Auto Post": "Agent บันทึกรายจ่ายอัตโนมัติ",
        "Agent Quotation": "Agent สร้างใบเสนอราคา",
        "Agent Contact": "Agent สร้างรายชื่อติดต่อ",
        "Agent Invoice": "Agent สร้างใบแจ้งหนี้",
        "live browser capture": "วิดีโออัดจาก browser จริง",
        "real-browser-recording": "วิดีโออัดจาก browser จริง",
    }
    result = text or ""
    for src, dst in replacements.items():
        result = result.replace(src, dst)
    result = result.replace("ไม่พบ การ", "ไม่พบข้อมูลการ")
    result = result.replace("ไม่พบ บิล", "ไม่พบบิล")
    result = re.sub(r"(.+?) \(\1\)", r"\1", result)
    return result


def display_title(item: dict) -> str:
    title = clean_text((item.get("title") or "").strip())
    if title:
        return title
    fallback = {
        "/purchases/orders/1": "รายละเอียดใบสั่งซื้อ",
        "/expenses/1": "รายละเอียดรายจ่าย",
    }
    return fallback.get(item.get("route", ""), item.get("route", "ไม่พบชื่อหน้า"))


def build_docx(data: dict) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("คู่มือการใช้งาน ERPTH Qacc")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(8, 61, 71)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("ภาษาไทยสำหรับผู้ใช้งาน พร้อมภาพหน้าจอจริงและตารางตรวจครบทุกหน้า")
    sr.font.name = "Arial"
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(82, 103, 113)

    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = False
    rows = [
        ("วันที่จัดทำ", data.get("runDate", "")),
        ("ระบบที่ใช้บันทึกหน้าจอ", data.get("baseUrl", "")),
        ("รหัสนำหน้าข้อมูลทดสอบ", data.get("demoDataPrefix", "")),
        ("หมายเหตุ", "สถานะของแต่ละหน้าถูกบันทึกจากระบบจริง ถ้า endpoint หรือสิทธิ์ไม่พร้อมจะแสดงตามจริงในตารางตรวจสอบ"),
    ]
    for i, row in enumerate(rows):
        set_cell_text(meta.rows[i].cells[0], row[0], True, "083D47")
        set_cell_text(meta.rows[i].cells[1], clean_text(row[1]))
        set_cell_shading(meta.rows[i].cells[0], "E9F7F8")

    doc.add_page_break()
    add_heading(doc, "วิธีใช้คู่มือนี้", 1)
    add_body(
        doc,
        "คู่มือนี้จัดเรียงตามงานจริงของผู้ใช้งาน ERPTH Qacc ไม่ได้เรียงตาม route ทางเทคนิคเพียงอย่างเดียว "
        "ทุก route ที่พบในแอป React ถูกตรวจและบันทึกไว้ในภาคผนวกท้ายเล่ม ภาพทั้งหมดมาจาก React ที่เชื่อมต่อกับ Odoo database q01 ผ่าน Vite proxy",
    )
    add_bullets(
        doc,
        [
            "ดูบทหลักเพื่อเรียนรู้งานประจำวันและลำดับการทำงาน",
            "ดูภาพหน้าจอประกอบเพื่อเทียบกับหน้าจอจริง",
            "ดูตาราง Route Coverage Matrix เมื่อต้องการตรวจว่าหน้าใดเปิดได้, ข้อมูลว่าง, ถูกจำกัดสิทธิ์ หรือแสดงข้อผิดพลาดจาก backend",
            "ดูคลิปวิดีโอในโฟลเดอร์ videos เพื่อเรียนรู้ขั้นตอนสำคัญ 6 ชุด",
        ],
    )

    chapters: dict[str, list[dict]] = {}
    for item in data["coverage"]:
        chapters.setdefault(item["chapter"], []).append(item)

    for chapter, items in chapters.items():
        doc.add_page_break()
        add_heading(doc, chapter, 1)
        add_bullets(doc, CHAPTER_GUIDANCE.get(chapter, []))
        for item in items:
            add_heading(doc, display_title(item), 2)
            add_body(doc, f"Route: {item['route']} | สถานะ: {status_th(item.get('status'), item.get('error'))}")
            if item.get("textSample"):
                add_body(doc, f"ข้อความที่พบในหน้า: {clean_text(item['textSample'])}")
            notes = []
            if item.get("dynamic"):
                notes.append("หน้านี้เป็น route แบบ dynamic ต้องใช้เลขอ้างอิงรายการหรือ template ที่มีอยู่จริงในการใช้งาน")
            if item.get("hasForm"):
                notes.append("มีช่องกรอกหรือฟอร์มสำหรับค้นหา/สร้าง/แก้ไขข้อมูล")
            if item.get("hasTable"):
                notes.append("มีตารางหรือรายการข้อมูลสำหรับตรวจสอบและเปิดรายละเอียด")
            if item.get("hasHorizontalOverflow"):
                notes.append("พบหน้ากว้างล้นแนวนอนระหว่างบันทึกหน้าจอ ควรตรวจ layout ก่อนเผยแพร่ production")
            if item.get("error"):
                notes.append(f"ข้อผิดพลาดระหว่างเปิดหน้า: {item['error']}")
            if notes:
                add_bullets(doc, notes)
            add_image(doc, ROOT / item["screenshot"], f"ภาพหน้าจอ: {display_title(item)}")

    doc.add_page_break()
    add_heading(doc, "วิดีโอการใช้งาน 6 คลิป", 1)
    video_table = doc.add_table(rows=1, cols=3)
    for i, label in enumerate(["คลิป", "ประเภท", "ไฟล์"]):
        set_cell_text(video_table.rows[0].cells[i], label, True, "FFFFFF")
        set_cell_shading(video_table.rows[0].cells[i], "083D47")
    for video in data.get("videoResults", []):
        cells = video_table.add_row().cells
        set_cell_text(cells[0], clean_text(video["title"]))
        set_cell_text(cells[1], clean_text(video.get("type", video.get("kind", "live browser capture"))))
        set_cell_text(cells[2], video["path"])

    doc.add_page_break()
    add_heading(doc, "Route Coverage Matrix", 1)
    table = doc.add_table(rows=1, cols=5)
    widths = [2.2, 3.4, 5.2, 2.2, 4.2]
    headers = ["บท", "หน้า", "Route", "สถานะ", "หมายเหตุ"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True, "FFFFFF")
        set_cell_shading(table.rows[0].cells[i], "083D47")
    for item in data["coverage"]:
        cells = table.add_row().cells
        note = "dynamic route" if item.get("dynamic") else ""
        if item.get("error"):
            note = (note + " " + item["error"]).strip()
        set_cell_text(cells[0], item["chapter"])
        set_cell_text(cells[1], display_title(item))
        set_cell_text(cells[2], item["route"])
        set_cell_text(cells[3], status_th(item.get("status"), item.get("error")))
        set_cell_text(cells[4], clean_text(note or item.get("textSample", "")[:90]))

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Cm(width)

    doc.save(DOCX_OUT)


def img_size(path: Path, max_w: float, max_h: float) -> tuple[float, float]:
    from PIL import Image as PILImage

    with PILImage.open(path) as img:
        w, h = img.size
    scale = min(max_w / w, max_h / h)
    return w * scale, h * scale


def build_pdf(data: dict) -> None:
    thai_font_path = Path("/Library/Fonts/Arial Unicode.ttf")
    if thai_font_path.exists():
        pdfmetrics.registerFont(TTFont("ThaiManual", str(thai_font_path)))
        pdfmetrics.registerFont(TTFont("ThaiManualBold", str(thai_font_path)))
        regular_font = "ThaiManual"
        bold_font = "ThaiManualBold"
    else:
        regular_font = "Helvetica"
        bold_font = "Helvetica-Bold"
    pdf = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        rightMargin=1.35 * cm,
        leftMargin=1.35 * cm,
        topMargin=1.35 * cm,
        bottomMargin=1.35 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ThaiTitle", parent=styles["Title"], fontName=bold_font, fontSize=22, textColor=colors.HexColor("#083D47"))
    h1 = ParagraphStyle("ThaiH1", parent=styles["Heading1"], fontName=bold_font, fontSize=15, textColor=colors.HexColor("#083D47"), spaceAfter=8)
    h2 = ParagraphStyle("ThaiH2", parent=styles["Heading2"], fontName=bold_font, fontSize=11, textColor=colors.HexColor("#1C5B67"), spaceAfter=5)
    body = ParagraphStyle("ThaiBody", parent=styles["BodyText"], fontName=regular_font, fontSize=9.5, leading=13, spaceAfter=5)
    small = ParagraphStyle("Small", parent=body, fontSize=8, leading=10, textColor=colors.HexColor("#526771"))

    story = [
        Paragraph("คู่มือการใช้งาน ERPTH Qacc", title_style),
        Paragraph("คู่มือภาษาไทยพร้อมภาพหน้าจอจริง ตารางตรวจครบทุกหน้า และวิดีโออัดการใช้งานจริง 6 คลิป", body),
        Spacer(1, 8),
        Paragraph(f"วันที่จัดทำ: {data.get('runDate', '')}", small),
        Paragraph(f"ระบบที่ใช้บันทึกหน้าจอ: {data.get('baseUrl', '')}", small),
        Paragraph(f"รหัสนำหน้าข้อมูลทดสอบ: {data.get('demoDataPrefix', '')}", small),
        PageBreak(),
    ]

    chapters: dict[str, list[dict]] = {}
    for item in data["coverage"]:
        chapters.setdefault(item["chapter"], []).append(item)

    for chapter, items in chapters.items():
        story.append(Paragraph(chapter, h1))
        for tip in CHAPTER_GUIDANCE.get(chapter, []):
            story.append(Paragraph(f"- {tip}", body))
        for item in items:
            chunks = [
                Paragraph(display_title(item), h2),
                Paragraph(f"Route: {item['route']} | สถานะ: {status_th(item.get('status'), item.get('error'))}", small),
            ]
            if item.get("textSample"):
                chunks.append(Paragraph(clean_text(item["textSample"]), body))
            img = ROOT / item["screenshot"]
            if img.exists():
                w, h = img_size(img, 17.5 * cm, 10.2 * cm)
                chunks.append(Image(str(img), width=w, height=h))
                chunks.append(Paragraph(f"ภาพหน้าจอ: {display_title(item)}", small))
            story.append(KeepTogether(chunks))
            story.append(Spacer(1, 8))
        story.append(PageBreak())

    story.append(Paragraph("วิดีโอการใช้งาน 6 คลิป", h1))
    video_rows = [["คลิป", "ประเภท", "ไฟล์"]]
    for video in data.get("videoResults", []):
        video_rows.append([clean_text(video["title"]), clean_text(video.get("type", video.get("kind", "live browser capture"))), video["path"]])
    vt = Table(video_rows, colWidths=[7.5 * cm, 2 * cm, 7.5 * cm])
    vt.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#083D47")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8CCD2")),
        ("FONT", (0, 0), (-1, -1), regular_font, 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(vt)
    story.append(PageBreak())

    story.append(Paragraph("Route Coverage Matrix", h1))
    rows = [["บท", "หน้า", "Route", "สถานะ"]]
    for item in data["coverage"]:
        rows.append([item["chapter"], display_title(item), item["route"], status_th(item.get("status"), item.get("error"))])
    matrix = Table(rows, repeatRows=1, colWidths=[3.2 * cm, 4.4 * cm, 6.7 * cm, 3.1 * cm])
    matrix.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#083D47")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8CCD2")),
        ("FONT", (0, 0), (-1, -1), regular_font, 6.8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(matrix)
    pdf.build(story)


def build_qa_report(data: dict) -> None:
    errors = [item for item in data["coverage"] if item.get("error") or item.get("status") == "error"]
    restricted = [item for item in data["coverage"] if item.get("status") == "restricted"]
    empty = [item for item in data["coverage"] if item.get("status") == "empty-state"]
    overflow = [item for item in data["coverage"] if item.get("hasHorizontalOverflow")]
    lines = [
        "# ERPTH Qacc User Manual QA Report",
        "",
        f"- Generated: {datetime.now().isoformat(timespec='seconds')}",
        f"- Capture base URL: {data.get('baseUrl', '')}",
        f"- Desktop screenshots: {len(data.get('coverage', []))}",
        "- Mobile version: not included per user request.",
        f"- Live capture videos: {len(data.get('videoResults', []))}",
        "- Video type: real browser recordings captured during navigation and converted to MP4.",
        "- DOCX render QA: skipped because `soffice` is not available in this environment.",
        "- PDF generated directly with ReportLab from the same capture result.",
        "",
        "## Status Summary",
        "",
        f"- Error pages: {len(errors)}",
        f"- Restricted pages: {len(restricted)}",
        f"- Empty-state pages: {len(empty)}",
        f"- Horizontal overflow flags: {len(overflow)}",
        "",
        "## Files",
        "",
        f"- DOCX: `{DOCX_OUT}`",
        f"- PDF: `{PDF_OUT}`",
        "- Videos: `Users Manaul Qacc/videos/*.mp4` live browser captures",
        "- Screenshots: `Users Manaul Qacc/screenshots/*.png`",
        "",
        "## Pages Needing Attention",
        "",
    ]
    attention = errors + restricted + overflow
    if not attention:
        lines.append("- None flagged.")
    else:
        for item in attention:
            lines.append(f"- `{item['route']}` - {display_title(item)} - {status_th(item.get('status'), item.get('error'))}")
    QA_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    if not DATA.exists():
        raise FileNotFoundError(f"Missing capture results: {DATA}")
    data = json.loads(DATA.read_text(encoding="utf-8"))
    if VIDEO_DATA.exists():
        video_data = json.loads(VIDEO_DATA.read_text(encoding="utf-8"))
        data["videoResults"] = video_data.get("results", [])
    else:
        data["videoResults"] = data.get("videoResults", [])
    data["mobileShots"] = []
    build_docx(data)
    build_pdf(data)
    build_qa_report(data)
    print(json.dumps({
        "docx": str(DOCX_OUT),
        "pdf": str(PDF_OUT),
        "qa": str(QA_OUT),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
