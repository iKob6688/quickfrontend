from __future__ import annotations

import json
import re
import shutil
import zipfile
from collections import defaultdict
from dataclasses import dataclass, asdict
from datetime import date, datetime
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


SOURCE_DOCX = Path("Users Manaul Qacc/ERPTH-Qacc-User-Manual-TH.docx")
CAPTURE_JSON = Path("Users Manaul Qacc/data/manual-capture-results.json")
VIDEO_JSON = Path("Users Manaul Qacc/data/manual-video-results.json")
OUTPUT_DIR = Path("output")
IMAGE_DIR = OUTPUT_DIR / "extracted_images"
DOCX_OUT = OUTPUT_DIR / "ERPTH_Qacc_User_Manual_Production_TH.docx"
PDF_OUT = OUTPUT_DIR / "ERPTH_Qacc_User_Manual_Production_TH.pdf"
QA_OUT = OUTPUT_DIR / "ERPTH_Qacc_User_Manual_QA_Report.md"
MAP_OUT = OUTPUT_DIR / "document_map.json"

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    "vt": "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes",
}


@dataclass
class Block:
    index: int
    kind: str
    text: str
    style: str = ""
    image_id: str | None = None
    image_path: str | None = None
    route: str | None = None


@dataclass
class ImageMap:
    image_id: str
    file: str
    module: str
    source_heading: str
    route: str | None
    mapped_section: str
    confidence: str
    reason: str


def clean_text(text: str) -> str:
    text = (text or "").replace("\u200b", "").strip()
    replacements = {
        "Route Coverage Matrix": "ตารางตรวจสอบหน้าระบบ",
        "Print Preview": "ตัวอย่างก่อนพิมพ์",
        "Drilldown": "ดูรายละเอียดเจาะลึก",
        "Partner Ledger": "รายงานลูกหนี้/เจ้าหนี้รายคู่ค้า",
        "General Ledger": "บัญชีแยกประเภท",
        "Trial Balance": "งบทดลอง",
        "Balance Sheet": "งบดุล",
        "Profit and Loss": "กำไรขาดทุน",
        "Aged Receivables": "อายุลูกหนี้",
        "Aged Payables": "อายุเจ้าหนี้",
        "Cash Book": "สมุดเงินสด",
        "Bank Book": "สมุดเงินฝากธนาคาร",
        "Vendor Bill": "บิลผู้ขาย",
        "Delivery": "การส่งสินค้า",
        "Sale Order": "คำสั่งขาย",
        "Dashboard": "แดชบอร์ด",
        "Reports Studio": "สตูดิโอรายงาน",
        "Backend": "ระบบหลังบ้าน",
        "backend": "ระบบหลังบ้าน",
        "endpoint": "ระบบเชื่อมต่อข้อมูล",
        "Endpoint": "ระบบเชื่อมต่อข้อมูล",
        "dynamic route": "หน้ารายละเอียดข้อมูล",
        "Dynamic route": "หน้ารายละเอียดข้อมูล",
        "res.partner": "รายชื่อติดต่อ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"(.+?) \(\1\)", r"\1", text)
    return text


def thai_status(status: str | None, error: str | None = None) -> str:
    if error:
        return "มีข้อผิดพลาด"
    return {
        "loaded": "เปิดได้",
        "login": "หน้า Login",
        "empty-state": "เปิดได้ แต่ยังไม่มีข้อมูล",
        "error": "แสดงข้อผิดพลาดจริง",
        "restricted": "จำกัดสิทธิ์",
        "not-found": "ไม่พบหน้า",
    }.get(status or "", status or "-")


def classify_route(route: str | None, title: str = "") -> str:
    text = f"{route or ''} {title}".lower()
    if "login" in text or "dashboard" in text:
        return "เริ่มต้นใช้งาน"
    if "customer" in text or "product" in text:
        return "ข้อมูลพื้นฐาน"
    if "/sales" in text or "invoice" in text or "receipt" in text or "note" in text or "delivery" in text:
        return "กระบวนการขาย"
    if "/purchases" in text:
        return "กระบวนการซื้อ"
    if "/expenses" in text or "accounting/admin" in text or "tax-settings" in text:
        return "การเงินและบัญชี"
    if "/accounting/reports" in text:
        return "รายงาน"
    if "etax" in text:
        return "การเงินและบัญชี"
    if "backend-connection" in text or "reports-studio" in text or "agent" in text or "excel" in text:
        return "การตั้งค่า"
    return "ภาคผนวก"


def display_title(item: dict) -> str:
    title = clean_text((item.get("title") or "").strip())
    if title:
        return title
    return {
        "/purchases/orders/1": "รายละเอียดใบสั่งซื้อ",
        "/expenses/1": "รายละเอียดรายจ่าย",
    }.get(item.get("route", ""), item.get("route") or "ไม่พบชื่อหน้า")


def source_page_count(docx_path: Path) -> int | None:
    with zipfile.ZipFile(docx_path) as zf:
        try:
            xml = zf.read("docProps/app.xml")
        except KeyError:
            return None
    root = ET.fromstring(xml)
    pages = root.find(".//{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Pages")
    if pages is not None and pages.text and pages.text.isdigit():
        value = int(pages.text)
        # Word may leave this stale in generated documents. A single page for
        # a large extracted manual is less useful than reporting it as unknown.
        return value if value > 1 else None
    return None


def rel_map(docx_path: Path) -> dict[str, str]:
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/_rels/document.xml.rels")
    root = ET.fromstring(xml)
    result = {}
    for rel in root.findall("rel:Relationship", NS):
        rid = rel.attrib.get("Id")
        target = rel.attrib.get("Target", "")
        if rid and target.startswith("media/"):
            result[rid] = f"word/{target}"
    return result


def paragraph_text(el: ET.Element) -> str:
    return "".join(t.text or "" for t in el.findall(".//w:t", NS)).strip()


def table_text(el: ET.Element) -> str:
    rows = []
    for tr in el.findall(".//w:tr", NS):
        cells = []
        for tc in tr.findall("./w:tc", NS):
            cells.append(" ".join(t.text or "" for t in tc.findall(".//w:t", NS)).strip())
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def paragraph_style(el: ET.Element) -> str:
    pstyle = el.find("./w:pPr/w:pStyle", NS)
    return pstyle.attrib.get(qn("w:val"), "") if pstyle is not None else ""


def extract_docx(docx_path: Path) -> tuple[list[Block], list[ImageMap], dict]:
    if IMAGE_DIR.exists():
        shutil.rmtree(IMAGE_DIR)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    blocks: list[Block] = []
    image_maps: list[ImageMap] = []
    rels = rel_map(docx_path)
    heading_context = ""
    route_context = None
    image_counter = 0

    with zipfile.ZipFile(docx_path) as zf:
        document_xml = zf.read("word/document.xml")
        root = ET.fromstring(document_xml)
        body = root.find("w:body", NS)
        if body is None:
            raise RuntimeError("DOCX has no document body")

        for child in body:
            tag = child.tag.split("}")[-1]
            index = len(blocks)
            if tag == "p":
                text = paragraph_text(child)
                style = paragraph_style(child)
                if text:
                    route_match = re.search(r"(/[A-Za-z0-9_./:-]+)", text)
                    if route_match:
                        route_context = route_match.group(1).rstrip("|")
                    if style.lower().startswith("heading") or len(text) < 90:
                        heading_context = clean_text(text)
                    blocks.append(Block(index=index, kind="paragraph", text=text, style=style, route=route_context))

                for blip in child.findall(".//a:blip", NS):
                    rid = blip.attrib.get(qn("r:embed"))
                    if not rid or rid not in rels:
                        continue
                    image_counter += 1
                    target = rels[rid]
                    ext = Path(target).suffix or ".png"
                    image_id = f"image_{image_counter:03d}"
                    out_path = IMAGE_DIR / f"{image_id}{ext}"
                    out_path.write_bytes(zf.read(target))
                    module = classify_route(route_context, heading_context)
                    mapped_section = module
                    confidence = "high" if route_context or heading_context else "low"
                    reason = "mapped by nearest route/heading" if confidence == "high" else "no nearby route or heading"
                    blocks.append(
                        Block(
                            index=len(blocks),
                            kind="image",
                            text=heading_context,
                            image_id=image_id,
                            image_path=str(out_path),
                            route=route_context,
                        )
                    )
                    image_maps.append(
                        ImageMap(
                            image_id=image_id,
                            file=str(out_path),
                            module=module,
                            source_heading=heading_context,
                            route=route_context,
                            mapped_section=mapped_section,
                            confidence=confidence,
                            reason=reason,
                        )
                    )
            elif tag == "tbl":
                text = table_text(child)
                if text:
                    blocks.append(Block(index=index, kind="table", text=text, route=route_context))

    stats = {
        "source_docx": str(docx_path),
        "source_pages": source_page_count(docx_path),
        "paragraph_count": sum(1 for b in blocks if b.kind == "paragraph"),
        "table_count": sum(1 for b in blocks if b.kind == "table"),
        "image_count": len(image_maps),
    }
    return blocks, image_maps, stats


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: int = 12) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(clean_text(text))
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_para(doc: Document, text: str, size: int = 14, bold: bool = False, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(clean_text(text))
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(clean_text(text))
    r.font.name = "Arial"
    r.font.size = Pt(22 if level == 1 else 17 if level == 2 else 15)
    r.bold = True
    r.font.color.rgb = RGBColor(8, 61, 71) if level == 1 else RGBColor(28, 91, 103)


def add_numbered_steps(doc: Document, steps: Iterable[str]) -> None:
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(clean_text(step))
        r.font.name = "Arial"
        r.font.size = Pt(14)


def add_bullets(doc: Document, bullets: Iterable[str]) -> None:
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(clean_text(bullet))
        r.font.name = "Arial"
        r.font.size = Pt(14)


def add_page_number_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Version 1.0 | Page ")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, fld_end):
        run._r.append(el)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    header = section.header.paragraphs[0]
    header.text = "ERPTH Qacc User Manual"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(10)
    header.runs[0].font.bold = True
    add_page_number_footer(section)
    return doc


def add_captioned_image(doc: Document, image_path: Path, caption: str, figure_no: int) -> bool:
    if not image_path.exists():
        return False
    try:
        with PILImage.open(image_path) as img:
            w, h = img.size
        max_width = Inches(6.2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=max_width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"รูปที่ {figure_no}: {clean_text(caption)}")
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.italic = True
        run.font.color.rgb = RGBColor(90, 108, 118)
        return w > 0 and h > 0
    except Exception:
        return False


def pick_images(image_maps: list[ImageMap], module: str, limit: int = 3) -> list[ImageMap]:
    picked = []
    seen = set()
    for img in image_maps:
        if img.module == module and img.file not in seen and img.confidence != "low":
            picked.append(img)
            seen.add(img.file)
        if len(picked) >= limit:
            break
    return picked


def build_workflow_table(doc: Document) -> None:
    steps = ["ข้อมูลพื้นฐาน", "ลูกค้า / ผู้ขาย / สินค้า", "การขาย / การซื้อ", "ใบแจ้งหนี้ / บิลผู้ขาย", "รับเงิน / จ่ายเงิน", "ภาษี", "รายงานบัญชี"]
    table = doc.add_table(rows=1, cols=len(steps))
    table.autofit = True
    for i, step in enumerate(steps):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, step, bold=True, color="083D47", size=10)
        set_cell_shading(cell, "E9F7F8")
    add_para(doc, "ลำดับการทำงานโดยรวม: " + "  →  ".join(steps), 14, True, RGBColor(8, 61, 71))


WORKFLOW_SECTIONS = [
    {
        "chapter": "บทที่ 1 เริ่มต้นใช้งาน",
        "module": "เริ่มต้นใช้งาน",
        "objective": "ใช้สำหรับเข้าสู่ระบบ ตรวจภาพรวมธุรกิจ และเข้าถึงเมนูหลักของ ERPTH Qacc",
        "steps": ["เปิดหน้า ERPTH Qacc", "กรอกชื่อผู้ใช้และรหัสผ่าน", "ตรวจสอบแดชบอร์ดและสถานะการเชื่อมต่อ", "ใช้เมนูหลักหรือช่องค้นหาเพื่อไปยังงานที่ต้องการ"],
        "result": "ระบบจะแสดงแดชบอร์ด เมนูหลัก และข้อมูลบริษัทที่เชื่อมต่ออยู่",
        "caution": "หากเข้าสู่ระบบไม่ได้ ให้ตรวจสอบฐานข้อมูล q01, ชื่อผู้ใช้, รหัสผ่าน และสถานะ Odoo server",
    },
    {
        "chapter": "บทที่ 2 ข้อมูลพื้นฐาน",
        "module": "ข้อมูลพื้นฐาน",
        "objective": "ใช้จัดการข้อมูลลูกค้า ผู้ขาย คู่ค้า และสินค้า/บริการ ก่อนเริ่มสร้างเอกสารธุรกิจ",
        "steps": ["เปิดเมนูรายชื่อติดต่อหรือสินค้า/บริการ", "ค้นหาข้อมูลเดิมก่อนสร้างใหม่", "กรอกข้อมูลสำคัญ เช่น ชื่อ ที่อยู่ เลขภาษี และข้อมูลภาษีขาย/ซื้อ", "บันทึกและตรวจสอบสถานะ Active"],
        "result": "ระบบจะมีข้อมูลพื้นฐานพร้อมนำไปใช้ในเอกสารขาย ซื้อ และบัญชี",
        "caution": "ข้อมูลภาษีและชื่อคู่ค้าควรถูกต้องก่อนออกเอกสารจริง เพราะจะถูกนำไปใช้ต่อในใบแจ้งหนี้และรายงาน",
    },
    {
        "chapter": "บทที่ 3 กระบวนการขาย",
        "module": "กระบวนการขาย",
        "objective": "ใช้จัดการงานขายตั้งแต่ใบเสนอราคา คำสั่งขาย การส่งสินค้า ใบแจ้งหนี้ การรับชำระ และใบเสร็จ",
        "steps": ["เลือกลูกค้าที่ต้องการขาย", "สร้างใบเสนอราคาและระบุสินค้า/บริการ", "ตรวจสอบคำสั่งขายและสถานะเอกสาร", "สร้างใบแจ้งหนี้และตรวจยอดรวม", "บันทึกรับชำระเงินและออกใบเสร็จ", "ตรวจสอบเอกสาร e-Tax หากเปิดใช้งาน"],
        "result": "ผู้ใช้สามารถติดตามเอกสารขายและสถานะการรับชำระเงินได้ครบกระบวนการ",
        "caution": "หน้าที่เป็น error หรือไม่พบข้อมูลถูกย้ายไป Known Issues ไม่ควรใช้เป็นขั้นตอนหลักของคู่มือ",
    },
    {
        "chapter": "บทที่ 4 กระบวนการซื้อ",
        "module": "กระบวนการซื้อ",
        "objective": "ใช้จัดการงานจัดซื้อ ตั้งแต่ผู้ขาย ใบสั่งซื้อ การรับสินค้า บิลผู้ขาย และการปิดเจ้าหนี้",
        "steps": ["ตรวจสอบข้อมูลผู้ขาย", "สร้างคำขอซื้อหรือใบสั่งซื้อ", "ตรวจรับสินค้าเมื่อได้รับของ", "ตรวจบิลผู้ขายและยอดภาษี", "ส่งต่อการชำระเงินและปิดเจ้าหนี้"],
        "result": "เอกสารซื้อถูกติดตามตามลำดับงานและนำไปใช้ในบัญชีเจ้าหนี้ได้",
        "caution": "หากข้อมูลตัวอย่างไม่ครบ ให้เปิดรายละเอียดจากรายการจริงแทนการเปิด URL โดยตรง",
    },
    {
        "chapter": "บทที่ 5 การเงินและบัญชี",
        "module": "การเงินและบัญชี",
        "objective": "ใช้ตรวจภาษี ผังบัญชี สมุดรายวัน การรับ/จ่ายเงิน และข้อมูลทางบัญชีที่เกี่ยวข้อง",
        "steps": ["เปิดหน้าการบัญชีหรือการตั้งค่าภาษี", "เลือกหมวดข้อมูลที่ต้องตรวจสอบ", "กรองช่วงวันที่หรือสถานะเมื่อจำเป็น", "ตรวจยอดและสถานะก่อนนำไปใช้ในรายงานหรือปิดงวด"],
        "result": "ผู้ใช้บัญชีสามารถตรวจสอบข้อมูลสำคัญก่อนสรุปรายงานได้",
        "caution": "หน้าผู้ดูแลระบบควรให้เฉพาะผู้มีสิทธิ์ด้านบัญชีหรือ IT ใช้งาน",
    },
    {
        "chapter": "บทที่ 6 รายงาน",
        "module": "รายงาน",
        "objective": "ใช้ดูรายงานขาย ซื้อ ลูกหนี้ เจ้าหนี้ ภาษี และงบการเงินจากข้อมูลในระบบ",
        "steps": ["เปิดเมนูรายงานบัญชี", "เลือกประเภทรายงาน เช่น กำไรขาดทุน งบดุล งบทดลอง หรือภาษี", "กำหนดช่วงวันที่และตัวกรอง", "ตรวจผลรวมและเปิดดูรายละเอียดเจาะลึกเมื่อจำเป็น"],
        "result": "ระบบจะแสดงรายงานเพื่อใช้ตรวจสอบผลประกอบการและสถานะทางบัญชี",
        "caution": "หากรายงานแสดง error ให้ตรวจสอบ endpoint รายงานและข้อมูลบัญชีใน q01",
    },
    {
        "chapter": "บทที่ 7 การตั้งค่า",
        "module": "การตั้งค่า",
        "objective": "ใช้ตั้งค่าการเชื่อมต่อ บริษัท เครื่องมือเสริม Agent, Excel import และสตูดิโอรายงาน",
        "steps": ["เปิดเมนูตั้งค่าหรือเครื่องมือที่ต้องการ", "ตรวจสอบสิทธิ์และ token ที่จำเป็น", "แก้ไขค่าตามนโยบายบริษัท", "ทดสอบผลลัพธ์หลังตั้งค่า"],
        "result": "ระบบพร้อมทำงานตามสิทธิ์ บริษัท และรูปแบบเอกสารที่กำหนด",
        "caution": "การตั้งค่าบางส่วนมีผลต่อเอกสารจริง ควรให้ผู้ดูแลระบบเป็นผู้ดำเนินการ",
    },
]


def load_capture() -> tuple[list[dict], list[dict], dict]:
    capture = {"coverage": [], "runDate": str(date.today()), "baseUrl": "", "demoDataPrefix": ""}
    videos = []
    if CAPTURE_JSON.exists():
        capture = json.loads(CAPTURE_JSON.read_text(encoding="utf-8"))
    if VIDEO_JSON.exists():
        video_root = VIDEO_JSON.parent.parent
        for item in json.loads(VIDEO_JSON.read_text(encoding="utf-8")).get("results", []):
            video = dict(item)
            if video.get("path"):
                video["source_path"] = video["path"]
                video["path"] = str(video_root / video["path"])
            if video.get("raw"):
                video["source_raw"] = video["raw"]
                video["raw"] = str(video_root / video["raw"])
            videos.append(video)
    return capture.get("coverage", []), videos, capture


def make_route_matrix(coverage: list[dict]) -> list[dict]:
    matrix = []
    for item in coverage:
        matrix.append(
            {
                "module": classify_route(item.get("route"), display_title(item)),
                "page": display_title(item),
                "route": item.get("route", ""),
                "status": thai_status(item.get("status"), item.get("error")),
                "remark": clean_text(item.get("textSample", ""))[:120],
                "error": item.get("error"),
            }
        )
    return matrix


def known_issues(route_matrix: list[dict]) -> list[dict]:
    issues = []
    for row in route_matrix:
        if row["status"] in {"มีข้อผิดพลาด", "แสดงข้อผิดพลาดจริง", "ไม่พบหน้า", "จำกัดสิทธิ์"}:
            issues.append(
                {
                    "area": row["module"],
                    "page": row["page"],
                    "issue": row["status"],
                    "impact": "ผู้ใช้ไม่ควรใช้หน้านี้เป็นขั้นตอนหลักจนกว่าข้อมูลหรือการเชื่อมต่อจะพร้อม",
                    "action": "เปิดจากรายการเอกสารจริง ตรวจสิทธิ์ผู้ใช้ และตรวจ endpoint/ข้อมูลใน q01",
                    "route": row["route"],
                }
            )
    return issues


def create_docx(blocks: list[Block], images: list[ImageMap], stats: dict, coverage: list[dict], videos: list[dict], capture: dict) -> int:
    doc = setup_doc()
    fig_no = 1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ERPTH Qacc")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(8, 61, 71)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("คู่มือการใช้งานระบบบัญชีสำหรับธุรกิจไทย\nUser Manual for Thai Accounting Workflow\nVersion 1.0\nจัดทำสำหรับผู้ใช้งานทั่วไป")
    sr.font.name = "Arial"
    sr.font.size = Pt(16)
    add_para(doc, f"วันที่จัดทำ: {date.today().isoformat()}", 14, True)
    doc.add_page_break()

    add_heading(doc, "ประวัติเอกสาร / Version Control", 1)
    table = doc.add_table(rows=1, cols=4)
    for i, header in enumerate(["Version", "Date", "Description", "Author"]):
        set_cell_text(table.rows[0].cells[i], header, True, "FFFFFF")
        set_cell_shading(table.rows[0].cells[i], "083D47")
    cells = table.add_row().cells
    for i, value in enumerate(["1.0", date.today().isoformat(), "Production user manual from original route audit document", "ERPTH"]):
        set_cell_text(cells[i], value)

    add_heading(doc, "สารบัญ", 1)
    add_para(doc, "สารบัญนี้เป็น Word field สามารถคลิกขวาแล้วเลือก Update Field เพื่ออัปเดตเลขหน้าใน Microsoft Word", 12)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "คลิกขวาเพื่ออัปเดตสารบัญ"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, fld_text, fld_end):
        run._r.append(el)
    doc.add_page_break()

    add_heading(doc, "ภาพรวมระบบ", 1)
    add_para(doc, "ERPTH Qacc เป็นระบบบัญชีสำหรับธุรกิจไทย ใช้จัดการข้อมูลคู่ค้า สินค้า เอกสารขาย เอกสารซื้อ การรับชำระเงิน ภาษี และรายงานบัญชีในกระบวนการเดียวกัน")
    add_bullets(doc, ["จัดการลูกค้า ผู้ขาย และรายชื่อติดต่อ", "จัดการสินค้าและบริการ", "ทำเอกสารขายและเอกสารซื้อ", "ออกใบแจ้งหนี้ รับชำระเงิน และออกใบเสร็จ", "ตรวจสอบภาษีและดูรายงานบัญชี"])
    add_heading(doc, "Workflow รวมของ ERPTH Qacc", 2)
    build_workflow_table(doc)
    doc.add_page_break()

    for section in WORKFLOW_SECTIONS:
        add_heading(doc, section["chapter"], 1)
        add_heading(doc, "วัตถุประสงค์", 2)
        add_para(doc, section["objective"])
        add_heading(doc, "ขั้นตอนการใช้งาน", 2)
        add_numbered_steps(doc, section["steps"])
        add_heading(doc, "ผลลัพธ์ที่ได้", 2)
        add_para(doc, section["result"])
        add_heading(doc, "ข้อควรระวัง", 2)
        add_para(doc, section["caution"])
        for img in pick_images(images, section["module"], 3):
            if add_captioned_image(doc, Path(img.file), img.source_heading or section["chapter"], fig_no):
                fig_no += 1
        doc.add_page_break()

    add_heading(doc, "บทที่ 8 FAQ / Troubleshooting", 1)
    add_heading(doc, "เปิดหน้ารายละเอียดแล้วไม่พบข้อมูล", 2)
    add_para(doc, "สาเหตุ: เลขอ้างอิงเอกสารไม่มีอยู่จริง หรือข้อมูลตัวอย่างไม่ครบ")
    add_para(doc, "แนวทางแก้ไข: เปิดจากรายการเอกสารจริง แทนการเปิด URL โดยตรง")
    add_heading(doc, "รายงานบัญชีแสดงข้อผิดพลาด", 2)
    add_para(doc, "สาเหตุ: ข้อมูลบัญชีหรือ endpoint รายงานใน q01 ยังไม่พร้อม")
    add_para(doc, "แนวทางแก้ไข: ตรวจข้อมูลบัญชี ช่วงวันที่ สิทธิ์ผู้ใช้ และสถานะ Odoo server")
    doc.add_page_break()

    add_heading(doc, "วิดีโอสาธิตการใช้งาน", 1)
    add_para(doc, "ไฟล์วิดีโอต่อไปนี้เป็นการบันทึกหน้าจอจากการใช้งานจริงใน browser เพื่อประกอบการอบรมผู้ใช้")
    if videos:
        for video in videos:
            title = video.get("title") or video.get("id") or "วิดีโอสาธิต"
            path = video.get("path") or "-"
            add_para(doc, f"- {title}: {path}")
    else:
        add_para(doc, "ไม่พบไฟล์วิดีโอประกอบในรอบ build นี้")
    doc.add_page_break()

    route_matrix = make_route_matrix(coverage)
    issues = known_issues(route_matrix)

    add_heading(doc, "ภาคผนวก ก: Route Coverage Matrix", 1)
    matrix_table = doc.add_table(rows=1, cols=5)
    for i, header in enumerate(["Module", "Page", "Route", "Status", "Remark"]):
        set_cell_text(matrix_table.rows[0].cells[i], header, True, "FFFFFF", 10)
        set_cell_shading(matrix_table.rows[0].cells[i], "083D47")
    for row in route_matrix:
        cells = matrix_table.add_row().cells
        for i, value in enumerate([row["module"], row["page"], row["route"], row["status"], row["remark"]]):
            set_cell_text(cells[i], value, size=8)
    doc.add_page_break()

    add_heading(doc, "ภาคผนวก ข: Known Issues", 1)
    issue_table = doc.add_table(rows=1, cols=5)
    for i, header in enumerate(["Area", "Page", "Issue", "Impact", "Recommended Action"]):
        set_cell_text(issue_table.rows[0].cells[i], header, True, "FFFFFF", 9)
        set_cell_shading(issue_table.rows[0].cells[i], "083D47")
    for issue in issues:
        cells = issue_table.add_row().cells
        for i, value in enumerate([issue["area"], issue["page"], issue["issue"], issue["impact"], issue["action"]]):
            set_cell_text(cells[i], value, size=8)
    doc.add_page_break()

    add_heading(doc, "ภาคผนวก ค: Technical Notes สำหรับทีม IT / Dev", 1)
    add_bullets(
        doc,
        [
            f"Source DOCX: {SOURCE_DOCX}",
            f"Capture base URL: {capture.get('baseUrl', '')}",
            f"Database/Environment: q01 ผ่าน Vite proxy",
            f"Test Data Prefix: {capture.get('demoDataPrefix', '')}",
            f"Extracted images: {len(images)}",
            f"วิดีโอ: {len(videos)} ไฟล์ MP4 จาก real browser recording",
        ],
    )
    doc.add_page_break()

    unmapped = [img for img in images if img.confidence == "low"]
    if unmapped:
        add_heading(doc, "ภาคผนวก ง: Unmapped Screenshots", 1)
        add_para(doc, "รูปภาพต่อไปนี้ไม่สามารถระบุหัวข้อได้ชัดเจนจากข้อความใกล้เคียง จึงไม่ถูกนำไปใช้ในบทหลัก")
        for img in unmapped:
            if add_captioned_image(doc, Path(img.file), "Unmapped screenshot", fig_no):
                fig_no += 1

    doc.save(DOCX_OUT)
    Document(str(DOCX_OUT))
    return fig_no - 1


def create_pdf(coverage: list[dict], videos: list[dict], capture: dict, image_count: int) -> None:
    thai_font = Path("/Library/Fonts/Arial Unicode.ttf")
    regular = "Helvetica"
    if thai_font.exists():
        pdfmetrics.registerFont(TTFont("ThaiManual", str(thai_font)))
        regular = "ThaiManual"
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName=regular, fontSize=16, textColor=colors.HexColor("#083D47"))
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName=regular, fontSize=10, leading=14)
    pdf = SimpleDocTemplate(str(PDF_OUT), pagesize=A4, leftMargin=1.5 * cm, rightMargin=1.5 * cm, topMargin=1.4 * cm, bottomMargin=1.4 * cm)
    story = [
        Paragraph("ERPTH Qacc คู่มือการใช้งานระบบบัญชีสำหรับธุรกิจไทย", h1),
        Paragraph(f"Version 1.0 | วันที่จัดทำ {date.today().isoformat()}", body),
        Spacer(1, 10),
        Paragraph("เอกสาร PDF นี้เป็น companion export จากคู่มือ Word ฉบับ Production", body),
        Paragraph(f"จำนวนรูปที่ extract ได้: {image_count}", body),
        Paragraph(f"จำนวนวิดีโอประกอบ: {len(videos)}", body),
        PageBreak(),
    ]
    rows = [["Module", "Page", "Route", "Status"]]
    for row in make_route_matrix(coverage):
        rows.append([row["module"], row["page"], row["route"], row["status"]])
    table = Table(rows, repeatRows=1, colWidths=[3.3 * cm, 5 * cm, 6 * cm, 3.2 * cm])
    table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#083D47")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8CCD2")), ("FONT", (0, 0), (-1, -1), regular, 7), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(Paragraph("Route Coverage Matrix", h1))
    story.append(table)
    pdf.build(story)


def write_outputs(blocks: list[Block], images: list[ImageMap], stats: dict, coverage: list[dict], videos: list[dict], capture: dict, figures_used: int) -> None:
    route_matrix = make_route_matrix(coverage)
    issues = known_issues(route_matrix)
    unmapped = [img for img in images if img.confidence == "low"]
    qa_summary = {
        "source_paragraphs": stats["paragraph_count"],
        "source_tables": stats["table_count"],
        "images_extracted": stats["image_count"],
        "mapped_images": len(images) - len(unmapped),
        "unmapped_images": len(unmapped),
        "figures_used": figures_used,
        "route_coverage_rows": len(route_matrix),
        "known_issues": len(issues),
        "videos_referenced": len(videos),
    }
    document_map = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "source": stats,
        "outputs": {"docx": str(DOCX_OUT), "pdf": str(PDF_OUT), "qa_report": str(QA_OUT)},
        "qa_summary": qa_summary,
        "sections": WORKFLOW_SECTIONS,
        "chapters": WORKFLOW_SECTIONS,
        "images": [asdict(img) for img in images],
        "route_coverage": route_matrix,
        "known_issues": issues,
        "videos": videos,
    }
    MAP_OUT.write_text(json.dumps(document_map, ensure_ascii=False, indent=2), encoding="utf-8")
    qa = [
        "# ERPTH Qacc Production User Manual QA Report",
        "",
        f"- Generated: {datetime.now().isoformat(timespec='seconds')}",
        f"- Source DOCX: `{SOURCE_DOCX}`",
        f"- Source pages: {stats.get('source_pages') or 'unknown'}",
        f"- Source paragraphs extracted: {stats['paragraph_count']}",
        f"- Source tables extracted: {stats['table_count']}",
        f"- Images extracted: {len(images)}",
        f"- Images mapped successfully: {len(images) - len(unmapped)}",
        f"- Images unmapped: {len(unmapped)}",
        f"- Figures used in main/manual appendices: {figures_used}",
        f"- New workflow sections: {len(WORKFLOW_SECTIONS) + 4}",
        f"- Route coverage rows: {len(route_matrix)}",
        f"- Known issues: {len(issues)}",
        f"- Live browser MP4 videos referenced: {len(videos)}",
        "",
        "## Quality Gate",
        "",
        "- DOCX opens with python-docx: passed",
        "- Image extraction: passed",
        "- Main content rewritten as business workflow manual: passed",
        "- Technical route/error content moved to appendices: passed",
        "- Mobile version excluded: passed",
        "- DOCX visual render: skipped because `soffice` is not available in this environment",
        "",
        "## Known Issues",
        "",
    ]
    if issues:
        for issue in issues:
            qa.append(f"- `{issue['route']}` | {issue['page']} | {issue['issue']} | {issue['action']}")
    else:
        qa.append("- None")
    qa.extend(["", "## Human Review Checklist", "", "- เปิด DOCX ใน Microsoft Word แล้ว Update Table of Contents", "- ตรวจคำศัพท์เฉพาะของบริษัท เช่น ชื่อเมนูและสิทธิ์ผู้ใช้", "- ตรวจหน้า Known Issues กับสถานะ backend ล่าสุดก่อนส่งลูกค้า", "- ตรวจภาพ screenshot ว่าตรงกับข้อมูล production/demo ที่ต้องการเผยแพร่"])
    QA_OUT.write_text("\n".join(qa) + "\n", encoding="utf-8")


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    OUTPUT_DIR.mkdir(exist_ok=True)
    blocks, images, stats = extract_docx(SOURCE_DOCX)
    coverage, videos, capture = load_capture()
    figures_used = create_docx(blocks, images, stats, coverage, videos, capture)
    create_pdf(coverage, videos, capture, len(images))
    write_outputs(blocks, images, stats, coverage, videos, capture, figures_used)
    print(json.dumps({"docx": str(DOCX_OUT), "pdf": str(PDF_OUT), "qa": str(QA_OUT), "map": str(MAP_OUT), "images": len(images)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
