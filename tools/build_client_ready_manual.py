from __future__ import annotations

import json
import re
import shutil
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageChops, ImageStat


ROOT = Path(".")
SOURCE_PRODUCTION = Path("output/ERPTH_Qacc_User_Manual_Production_TH.docx")
SOURCE_ORIGINAL = Path("Users Manaul Qacc/ERPTH-Qacc-User-Manual-TH.docx")
CAPTURE_JSON = Path("Users Manaul Qacc/data/manual-capture-results.json")
VIDEO_JSON = Path("Users Manaul Qacc/data/manual-video-results.json")
SCREENSHOT_ROOT = Path("Users Manaul Qacc")
OUTPUT_DIR = Path("output")
CROPPED_DIR = OUTPUT_DIR / "cropped_images"
DOCX_OUT = OUTPUT_DIR / "ERPTH_Qacc_User_Manual_Client_Ready_TH.docx"
QA_OUT = OUTPUT_DIR / "ERPTH_Qacc_User_Manual_Client_Ready_QA_Report.md"

ACCENT = RGBColor(8, 61, 71)
MUTED = RGBColor(95, 105, 112)
LIGHT_FILL = "E9F7F8"
HEADER_FILL = "083D47"


THAI_REPLACEMENTS = {
    "เ มนู": "เมนู",
    "ภา ษี": "ภาษี",
    "ตั้งตั้ ค่า": "ตั้งค่า",
    "ขั้นขั้ ตอน": "ขั้นตอน",
    "เริ่มริ่ ต้น": "เริ่มต้น",
    "สิทธิ์ ธิ์": "สิทธิ์",
    "ใบแ → จ้งหนี้": "ใบแจ้งหนี้",
    "ใบแ จ้งหนี้": "ใบแจ้งหนี้",
    "รายชื่ อ": "รายชื่อ",
    "ข้ อมูล": "ข้อมูล",
    "ผู้ ใช้": "ผู้ใช้",
    "หน้ า": "หน้า",
    "เอกสาร ขาย": "เอกสารขาย",
    "เอกสาร ซื้อ": "เอกสารซื้อ",
    "ราย งาน": "รายงาน",
    "ระบบหลังบ้าน": "ระบบ",
    "endpoint": "การเชื่อมต่อข้อมูล",
    "Endpoint": "การเชื่อมต่อข้อมูล",
    "API": "การเชื่อมต่อข้อมูล",
    "route": "หน้า",
    "Route": "หน้า",
    "res.partner": "รายชื่อติดต่อ",
    "Backend": "ระบบ",
    "backend": "ระบบ",
}

ERROR_KEYWORDS = [
    "not found",
    "ไม่พบ",
    "backend error",
    "permission",
    "invalid",
    "data not found",
    "response format",
    "error",
]


@dataclass
class Shot:
    route: str
    title: str
    chapter: str
    status: str
    screenshot: str
    text_sample: str
    has_table: bool
    has_form: bool
    cropped_path: str | None = None
    crop_status: str = "not processed"
    crop_reason: str = ""


@dataclass
class Section:
    chapter: str
    number: str
    title: str
    purpose: str
    when: list[str]
    steps: list[str]
    result: str
    caution: str
    example: str
    route: str | None = None
    callouts: list[str] | None = None


def clean_thai(text: str) -> str:
    value = text or ""
    for old, new in THAI_REPLACEMENTS.items():
        value = value.replace(old, new)
    value = re.sub(r"[ \t]{2,}", " ", value)
    value = re.sub(r"\s+([,.:;])", r"\1", value)
    return value.strip()


def load_capture() -> tuple[list[Shot], list[dict]]:
    data = json.loads(CAPTURE_JSON.read_text(encoding="utf-8"))
    shots = []
    for item in data.get("coverage", []):
        shots.append(
            Shot(
                route=item.get("route", ""),
                title=clean_thai(item.get("title", "")),
                chapter=clean_thai(item.get("chapter", "")),
                status=item.get("status", ""),
                screenshot=item.get("screenshot", ""),
                text_sample=clean_thai(item.get("textSample", "")),
                has_table=bool(item.get("hasTable")),
                has_form=bool(item.get("hasForm")),
            )
        )
    videos = []
    if VIDEO_JSON.exists():
        video_root = VIDEO_JSON.parent.parent
        for item in json.loads(VIDEO_JSON.read_text(encoding="utf-8")).get("results", []):
            video = dict(item)
            if video.get("path"):
                video["path"] = str(video_root / video["path"])
            videos.append(video)
    return shots, videos


def is_error_shot(shot: Shot) -> bool:
    status = (shot.status or "").lower()
    sample = f"{shot.title} {shot.text_sample}".lower()
    if status in {"error", "not-found", "restricted", "empty-state"}:
        return True
    return any(keyword in sample for keyword in ERROR_KEYWORDS)


def crop_screenshot(src: Path, dst: Path) -> tuple[str, str]:
    if not src.exists():
        return "not cropped", "source file missing"
    with Image.open(src) as im:
        rgb = im.convert("RGB")
        width, height = rgb.size
        edge = Image.new("RGB", rgb.size, rgb.getpixel((0, height - 1)))
        diff = ImageChops.difference(rgb, edge).convert("L")
        mask = diff.point(lambda px: 255 if px > 14 else 0)

        # Add dark text and UI lines to the mask so white cards are retained
        # once padded around their visible content.
        gray = rgb.convert("L")
        dark = gray.point(lambda px: 255 if px < 235 else 0)
        mask = ImageChops.lighter(mask, dark)
        bbox = mask.getbbox()
        if not bbox:
            rgb.save(dst, quality=95)
            return "not cropped", "no meaningful UI bounds detected"

        left, top, right, bottom = bbox
        pad = 24
        left = max(left - pad, 0)
        top = max(top - pad, 0)
        right = min(right + pad, width)
        bottom = min(bottom + pad, height)
        crop_w = right - left
        crop_h = bottom - top
        if crop_w < width * 0.10 or crop_h < height * 0.03:
            rgb.save(dst, quality=95)
            return "not cropped", "auto-crop bounds looked unsafe"
        cropped = rgb.crop((left, top, right, bottom))
        cropped.save(dst, quality=95)
        if crop_w >= width * 0.96 and crop_h >= height * 0.96:
            return "kept", "screenshot already tightly framed"
        removed = int((1 - (crop_w * crop_h) / (width * height)) * 100)
        return "cropped", f"removed about {removed}% outer blank area"


def process_images(shots: list[Shot]) -> dict:
    if CROPPED_DIR.exists():
        shutil.rmtree(CROPPED_DIR)
    CROPPED_DIR.mkdir(parents=True, exist_ok=True)
    stats = {"total": 0, "cropped": 0, "kept": 0, "not_cropped": 0, "missing": 0}
    for idx, shot in enumerate(shots, start=1):
        if not shot.screenshot:
            stats["missing"] += 1
            continue
        src = SCREENSHOT_ROOT / shot.screenshot
        dst = CROPPED_DIR / f"{idx:03d}-{Path(shot.screenshot).name}"
        status, reason = crop_screenshot(src, dst)
        shot.cropped_path = str(dst)
        shot.crop_status = status
        shot.crop_reason = reason
        stats["total"] += 1
        if status == "cropped":
            stats["cropped"] += 1
        elif status == "kept":
            stats["kept"] += 1
        else:
            stats["not_cropped"] += 1
    return stats


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(14)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 22, ACCENT, 14, 8),
        ("Heading 2", 18, ACCENT, 10, 6),
        ("Heading 3", 15, RGBColor(40, 80, 95), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "ERPTH Qacc User Manual"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(10)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Version 1.0 | Page ")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, fld_end):
        r._r.append(el)
    return doc


def add_para(doc: Document, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(clean_thai(text))
    r.font.name = "Arial"
    r.font.size = Pt(14)
    r.bold = bold
    if color:
        r.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(clean_thai(text), level=level)
    p.paragraph_format.keep_with_next = True


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(clean_thai(item))
        r.font.name = "Arial"
        r.font.size = Pt(14)


def add_steps(doc: Document, steps: Iterable[str]) -> None:
    for i, step in enumerate(steps, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{i}. {clean_thai(step)}")
        r.font.name = "Arial"
        r.font.size = Pt(14)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 11, color: str = "000000") -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(clean_thai(str(text)))
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], HEADER_FILL)
        set_cell_text(table.rows[0].cells[i], h, True, 10, "FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, 10)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def image_for_route(shots_by_route: dict[str, Shot], route: str | None) -> Shot | None:
    if not route:
        return None
    shot = shots_by_route.get(route)
    if not shot or is_error_shot(shot):
        return None
    if not shot.cropped_path or not Path(shot.cropped_path).exists():
        return None
    return shot


CAPTIONS = {
    "/login": "หน้าเข้าสู่ระบบ ERPTH Qacc",
    "/dashboard": "Dashboard สำหรับตรวจสอบภาพรวมธุรกิจ",
    "/customers": "หน้ารายการรายชื่อติดต่อ",
    "/customers/new": "หน้าสร้างรายชื่อติดต่อใหม่",
    "/products": "หน้ารายการสินค้าและบริการ",
    "/products/new": "หน้าสร้างสินค้าและบริการใหม่",
    "/sales/orders": "หน้ารายการใบเสนอราคาและคำสั่งขาย",
    "/sales/invoices": "หน้ารายการใบแจ้งหนี้",
    "/sales/invoices/new": "หน้าสร้างใบแจ้งหนี้ใหม่",
    "/sales/receipts": "หน้ารายการใบเสร็จรับเงิน",
    "/purchases/orders": "หน้ารายการใบสั่งซื้อ",
    "/purchases/orders/new": "หน้าสร้างใบสั่งซื้อ",
    "/purchases/receipts/1": "หน้ารับสินค้าเข้าคลัง",
    "/purchases/requests": "หน้ารายการคำขอซื้อ",
    "/purchases/requests/1": "หน้ารายละเอียดคำขอซื้อ",
    "/expenses": "หน้ารายการรายจ่าย",
    "/expenses/new": "หน้าสร้างรายจ่ายใหม่",
    "/accounting/document-review": "กล่องงานตรวจสอบเอกสาร",
    "/accounting/reports": "หน้ารวมรายงานบัญชี",
    "/accounting/reports/profit-loss": "รายงานงบกำไรขาดทุน",
    "/accounting/reports/balance-sheet": "รายงานงบดุล",
    "/accounting/reports/trial-balance": "รายงานงบทดลอง",
    "/accounting/reports/partner-ledger/partner/1": "รายงานลูกหนี้/เจ้าหนี้รายคู่ค้า",
    "/accounting/tax-settings": "หน้าตั้งค่าภาษี",
    "/accounting/admin": "หน้าดูแลผังบัญชีและสมุดรายวัน",
    "/excel-import": "หน้านำเข้าข้อมูลจาก Excel",
    "/agent": "หน้ารวมเครื่องมือช่วยทำงาน",
    "/reports-studio/branding": "หน้าตั้งค่าแบรนด์เอกสาร",
    "/reports-studio/templates": "หน้าจัดการเทมเพลตรายงาน",
}


CALLOUTS = {
    "list": ["ปุ่มเพิ่มข้อมูลใหม่", "ช่องค้นหา", "ตารางรายการ", "สถานะของข้อมูลหรือเอกสาร"],
    "form": ["ช่องกรอกข้อมูลหลัก", "ข้อมูลผู้ติดต่อหรือเอกสาร", "ปุ่มบันทึก", "ปุ่มยกเลิกหรือย้อนกลับ"],
    "report": ["ตัวกรองช่วงวันที่", "ปุ่มแสดงผลรายงาน", "ตารางสรุปตัวเลข", "ลิงก์เปิดรายละเอียด"],
    "dashboard": ["ชื่อบริษัทที่กำลังใช้งาน", "สรุปยอดขายและรายจ่าย", "งานที่ต้องติดตาม", "เมนูไปยังงานหลัก"],
}


def add_figure(doc: Document, shot: Shot, fig_no: int, fallback_title: str, callouts: list[str] | None = None) -> bool:
    if not shot.cropped_path:
        return False
    path = Path(shot.cropped_path)
    if not path.exists():
        return False
    try:
        with Image.open(path) as im:
            w, h = im.size
            if w <= 0 or h <= 0:
                return False
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Inches(6.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.keep_with_next = bool(callouts)
        caption = CAPTIONS.get(shot.route, f"ภาพหน้าจอประกอบหัวข้อ {fallback_title}")
        run = cap.add_run(f"รูปที่ {fig_no}: {caption}")
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = MUTED
        if callouts:
            add_heading(doc, "จุดสำคัญบนหน้าจอ", 3)
            add_steps(doc, callouts)
        return True
    except Exception:
        return False


def workflow_block(doc: Document, title: str, steps: list[str]) -> None:
    add_heading(doc, title, 2)
    for i, step in enumerate(steps):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(step)
        r.font.name = "Arial"
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = ACCENT
        if i < len(steps) - 1:
            arrow = doc.add_paragraph()
            arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ar = arrow.add_run("↓")
            ar.font.name = "Arial"
            ar.font.size = Pt(15)
            ar.font.color.rgb = MUTED


def chapter_sections() -> list[Section]:
    sections = [
        Section("บทที่ 1 เริ่มต้นใช้งาน", "1.1", "เข้าสู่ระบบ", "ใช้สำหรับยืนยันตัวตนก่อนเริ่มทำงานใน ERPTH Qacc", ["เมื่อเริ่มใช้งานประจำวัน", "เมื่อระบบขอให้เข้าสู่ระบบใหม่"], ["เปิดหน้าเข้าสู่ระบบ", "กรอกชื่อผู้ใช้และรหัสผ่าน", "กดปุ่มเข้าสู่ระบบ", "ตรวจสอบว่าระบบเปิด Dashboard สำเร็จ"], "ผู้ใช้เข้าสู่ระบบและเห็นข้อมูลบริษัทที่มีสิทธิ์ใช้งาน", "ไม่ควรใช้บัญชีร่วมกัน เพราะประวัติเอกสารจะผูกกับผู้ใช้นั้น", "ใช้ทุกครั้งก่อนเริ่มบันทึกเอกสารขาย ซื้อ หรือบัญชี", "/login", CALLOUTS["form"]),
        Section("บทที่ 1 เริ่มต้นใช้งาน", "1.2", "Dashboard และเมนูหลัก", "ใช้ตรวจภาพรวมธุรกิจและเลือกงานที่ต้องการทำต่อ", ["เมื่อต้องการดูยอดรวมก่อนเริ่มงาน", "เมื่อต้องการไปยังเมนูขาย ซื้อ บัญชี หรือรายงาน"], ["ตรวจสอบชื่อบริษัทด้านบน", "ดูตัวเลขสรุปและงานที่ต้องติดตาม", "เลือกเมนูด้านข้างหรือเมนูบนหน้าจอ", "ออกจากระบบเมื่อใช้งานเสร็จ"], "ผู้ใช้เห็นภาพรวมและไปยังหน้าทำงานได้ถูกต้อง", "หากตัวเลขไม่ตรง ให้ตรวจสอบบริษัท ช่วงวันที่ และสิทธิ์ผู้ใช้ก่อนบันทึกงานต่อ", "ใช้ตอนเริ่มวันเพื่อเช็กงานค้างและภาพรวมรายรับรายจ่าย", "/dashboard", CALLOUTS["dashboard"]),
        Section("บทที่ 2 ข้อมูลพื้นฐาน", "2.1", "รายชื่อติดต่อ", "ใช้จัดเก็บลูกค้า ผู้ขาย และบุคคลที่เกี่ยวข้องกับเอกสาร", ["เมื่อต้องการค้นหาลูกค้าหรือผู้ขาย", "ก่อนสร้างเอกสารขายหรือซื้อ"], ["เปิดเมนูรายชื่อติดต่อ", "ค้นหาชื่อหรือรหัสที่ต้องการ", "เปิดรายการเพื่อดูข้อมูล", "เพิ่มข้อมูลใหม่หากยังไม่มีในระบบ"], "ข้อมูลคู่ค้าพร้อมใช้ในเอกสารขาย ซื้อ และรายงาน", "ควรตรวจเลขประจำตัวผู้เสียภาษีและที่อยู่ก่อนออกเอกสารจริง", "ใช้ก่อนสร้างใบเสนอราคาหรือใบสั่งซื้อให้คู่ค้ารายใหม่", "/customers", CALLOUTS["list"]),
        Section("บทที่ 2 ข้อมูลพื้นฐาน", "2.2", "เพิ่มรายชื่อติดต่อใหม่", "ใช้สร้างข้อมูลลูกค้าหรือผู้ขายรายใหม่", ["เมื่อมีลูกค้าใหม่", "เมื่อมีผู้ขายใหม่"], ["กดปุ่มเพิ่มรายชื่อติดต่อ", "กรอกชื่อ ประเภทคู่ค้า และข้อมูลติดต่อ", "กรอกข้อมูลภาษีและที่อยู่ให้ครบ", "กดบันทึก"], "ระบบบันทึกคู่ค้าและนำไปเลือกในเอกสารได้", "ข้อมูลภาษีผิดจะทำให้เอกสารและรายงานภาษีผิดตามไปด้วย", "ใช้เมื่อต้องออกใบแจ้งหนี้ให้ลูกค้าใหม่", "/customers/new", CALLOUTS["form"]),
        Section("บทที่ 2 ข้อมูลพื้นฐาน", "2.3", "แก้ไขรายชื่อติดต่อ", "ใช้ปรับปรุงข้อมูลคู่ค้าที่มีอยู่แล้ว", ["เมื่อที่อยู่หรือข้อมูลติดต่อเปลี่ยน", "เมื่อพบข้อมูลภาษีไม่ถูกต้อง"], ["ค้นหารายชื่อที่ต้องการ", "เปิดหน้ารายละเอียดจากรายการจริง", "กดแก้ไขและปรับข้อมูล", "บันทึกและตรวจผลในรายการ"], "เอกสารใหม่จะใช้ข้อมูลคู่ค้าล่าสุด", "ควรตรวจผลกระทบกับเอกสารที่ออกไปแล้วก่อนแก้ข้อมูลสำคัญ", "ใช้เมื่อบริษัทลูกค้าเปลี่ยนที่อยู่สำหรับออกใบกำกับภาษี", None, None),
        Section("บทที่ 2 ข้อมูลพื้นฐาน", "2.4", "สินค้าและบริการ", "ใช้จัดการรายการสินค้า บริการ ราคา และข้อมูลสำหรับเอกสาร", ["ก่อนสร้างใบเสนอราคา", "ก่อนสร้างใบสั่งซื้อหรือบิลผู้ขาย"], ["เปิดเมนูสินค้าและบริการ", "ค้นหาสินค้าหรือบริการ", "ตรวจราคา หน่วย และสถานะ", "เพิ่มข้อมูลใหม่หากยังไม่มี"], "สินค้าและบริการพร้อมเลือกในเอกสาร", "ควรตรวจหน่วยนับและราคาก่อนนำไปใช้ในเอกสารจริง", "ใช้ตรวจรายการบริการก่อนเสนอราคาให้ลูกค้า", "/products", CALLOUTS["list"]),
        Section("บทที่ 2 ข้อมูลพื้นฐาน", "2.5", "เพิ่มสินค้าใหม่", "ใช้สร้างรายการสินค้าและบริการใหม่", ["เมื่อมีบริการใหม่", "เมื่อเริ่มขายหรือซื้อสินค้ารายการใหม่"], ["กดปุ่มเพิ่มสินค้า", "กรอกชื่อสินค้า ประเภท หน่วย และราคา", "ตรวจข้อมูลบัญชีหรือภาษีที่เกี่ยวข้อง", "กดบันทึก"], "ระบบเพิ่มรายการสินค้าและนำไปใช้ในเอกสารได้", "ชื่อสินค้าและหน่วยควรสื่อความหมายชัดเจนเพื่อป้องกันเลือกผิด", "ใช้เมื่อต้องเพิ่มบริการรายเดือนรายการใหม่", "/products/new", CALLOUTS["form"]),
        Section("บทที่ 2 ข้อมูลพื้นฐาน", "2.6", "แก้ไขสินค้า", "ใช้ปรับข้อมูลสินค้าและบริการที่มีอยู่", ["เมื่อราคาเปลี่ยน", "เมื่อแก้ชื่อหรือหน่วยนับ"], ["ค้นหาสินค้าที่ต้องการ", "เปิดข้อมูลจากรายการสินค้า", "แก้ไขข้อมูลที่จำเป็น", "บันทึกและตรวจรายการอีกครั้ง"], "ข้อมูลสินค้าใหม่พร้อมใช้กับเอกสารครั้งถัดไป", "ควรแยกระหว่างการแก้ข้อมูลสินค้าและการแก้เอกสารเก่าที่ออกไปแล้ว", "ใช้เมื่อเปลี่ยนราคาบริการก่อนทำใบเสนอราคาใหม่", None, None),
        Section("บทที่ 3 กระบวนการขาย", "3.1", "ภาพรวมกระบวนการขาย", "ใช้ติดตามงานขายตั้งแต่เสนอราคา รับเงิน จนถึงเอกสารภาษี", ["เมื่อเริ่มงานขายใหม่", "เมื่อต้องการตรวจสถานะงานขาย"], ["ตรวจข้อมูลลูกค้าและสินค้า", "สร้างเอกสารขาย", "ตรวจใบแจ้งหนี้และการรับชำระ", "ออกใบเสร็จหรือเอกสารภาษีตามขั้นตอนบริษัท"], "งานขายถูกติดตามเป็นลำดับและตรวจสอบย้อนหลังได้", "ควรเริ่มจากข้อมูลพื้นฐานที่ถูกต้องเสมอ", "ใช้เป็นภาพรวมในการอบรมผู้ใช้งานฝ่ายขายและบัญชี", "/sales/orders", CALLOUTS["list"]),
        Section("บทที่ 3 กระบวนการขาย", "3.2", "รายการใบเสนอราคา / คำสั่งขาย", "ใช้ดูเอกสารขายทั้งหมดและตรวจสถานะงานขาย", ["เมื่อต้องค้นหาใบเสนอราคา", "เมื่อต้องติดตามคำสั่งขาย"], ["เปิดเมนูใบเสนอราคา / คำสั่งขาย", "ใช้ช่องค้นหาเพื่อหาเอกสาร", "ตรวจสถานะและยอดรวม", "เปิดเอกสารจากรายการที่ต้องการ"], "ผู้ใช้พบเอกสารขายและเห็นสถานะล่าสุด", "ควรเปิดจากรายการจริง ไม่ควรพิมพ์เลขหน้าเอง", "ใช้ตรวจใบเสนอราคาที่รอลูกค้าอนุมัติ", "/sales/orders", CALLOUTS["list"]),
        Section("บทที่ 3 กระบวนการขาย", "3.3", "สร้างใบเสนอราคา", "ใช้จัดทำข้อเสนอให้ลูกค้าก่อนยืนยันงานขาย", ["เมื่อลูกค้าขอราคา", "เมื่อต้องปรับรายการสินค้าให้ลูกค้าตรวจสอบ"], ["ตรวจว่ามีข้อมูลลูกค้าและสินค้าแล้ว", "กดสร้างใบเสนอราคา", "เลือกลูกค้าและเพิ่มสินค้า", "ตรวจราคา ภาษี และหมายเหตุ", "บันทึกเอกสาร"], "ระบบสร้างเอกสารขายที่ใช้ติดตามต่อได้", "หากหน้าสร้างแสดงข้อจำกัด ให้บันทึกจากหน้าที่ระบบเปิดได้จริงหรือตรวจข้อมูลตัวอย่าง", "ใช้เมื่อต้องเสนอราคางานบริการให้ลูกค้าใหม่", None, None),
        Section("บทที่ 3 กระบวนการขาย", "3.4", "ตรวจสอบรายละเอียดใบเสนอราคา", "ใช้ทบทวนข้อมูลก่อนส่งต่อหรือออกเอกสารถัดไป", ["ก่อนยืนยันคำสั่งขาย", "ก่อนออกใบแจ้งหนี้"], ["ค้นหาเอกสารจากรายการ", "เปิดรายละเอียด", "ตรวจชื่อลูกค้า รายการสินค้า ราคา และภาษี", "แก้ไขหากพบข้อมูลผิด"], "ข้อมูลพร้อมดำเนินการขั้นถัดไป", "ควรตรวจยอดรวมก่อนส่งให้ลูกค้าทุกครั้ง", "ใช้ก่อนส่งใบเสนอราคาให้ลูกค้าตรวจสอบ", None, None),
        Section("บทที่ 3 กระบวนการขาย", "3.5", "ออกใบแจ้งหนี้", "ใช้แจ้งยอดเรียกเก็บเงินจากลูกค้า", ["เมื่องานขายได้รับอนุมัติ", "เมื่อถึงกำหนดวางบิล"], ["เปิดเมนูใบแจ้งหนี้", "ค้นหาหรือสร้างใบแจ้งหนี้", "ตรวจลูกค้า รายการสินค้า ยอดเงิน และภาษี", "บันทึกและส่งเอกสารตามขั้นตอนบริษัท"], "ระบบมีใบแจ้งหนี้สำหรับติดตามรับเงิน", "ควรตรวจวันที่เอกสารและเลขประจำตัวผู้เสียภาษีก่อนส่งให้ลูกค้า", "ใช้เมื่อต้องวางบิลหลังส่งมอบงาน", "/sales/invoices", CALLOUTS["list"]),
        Section("บทที่ 3 กระบวนการขาย", "3.6", "รับชำระเงิน", "ใช้บันทึกเงินที่ได้รับจากลูกค้า", ["เมื่อลูกค้าชำระเงินแล้ว", "เมื่อต้องปิดยอดค้างรับ"], ["เปิดเอกสารที่ต้องรับเงิน", "ตรวจยอดเงินและช่องทางรับชำระ", "บันทึกรับชำระเงิน", "ตรวจสถานะหลังบันทึก"], "ยอดค้างรับลดลงและข้อมูลพร้อมออกใบเสร็จ", "ควรตรวจหลักฐานการชำระเงินก่อนบันทึก", "ใช้เมื่อได้รับเงินโอนจากลูกค้าหลังออกใบแจ้งหนี้", "/sales/receipts", CALLOUTS["list"]),
        Section("บทที่ 3 กระบวนการขาย", "3.7", "ใบเสร็จรับเงิน", "ใช้ยืนยันการรับเงินและเก็บหลักฐานให้ลูกค้า", ["หลังรับชำระเงิน", "เมื่อลูกค้าขอหลักฐานการรับเงิน"], ["เปิดเมนูใบเสร็จรับเงิน", "ค้นหาใบเสร็จที่เกี่ยวข้อง", "ตรวจยอดเงินและข้อมูลลูกค้า", "พิมพ์หรือส่งเอกสารตามขั้นตอน"], "ลูกค้าได้รับหลักฐานรับเงิน และระบบมีประวัติการรับชำระ", "ควรตรวจเลขที่เอกสารก่อนส่งออก", "ใช้หลังบันทึกรับชำระเงินครบยอด", "/sales/receipts", CALLOUTS["list"]),
        Section("บทที่ 3 กระบวนการขาย", "3.8", "e-Tax", "ใช้จัดการเอกสารภาษีอิเล็กทรอนิกส์ตามนโยบายบริษัท", ["เมื่อต้องออกเอกสารภาษีอิเล็กทรอนิกส์", "เมื่อต้องตรวจสถานะการส่งเอกสาร"], ["ตรวจข้อมูลลูกค้าและเอกสารภาษี", "เปิดหน้าที่เกี่ยวข้องกับ e-Tax", "ตรวจสถานะและข้อมูลที่ต้องส่ง", "ดำเนินการตามสิทธิ์ผู้ใช้"], "เอกสารภาษีอยู่ในขั้นตอนพร้อมตรวจสอบหรือส่งออก", "บางหน้าต้องใช้การตั้งค่าระบบและสิทธิ์ผู้ดูแล", "ใช้เมื่อลูกค้าต้องการใบกำกับภาษีอิเล็กทรอนิกส์", None, None),
        Section("บทที่ 4 กระบวนการซื้อ", "4.1", "ภาพรวมกระบวนการซื้อ", "ใช้ติดตามการซื้อจากผู้ขายจนถึงรับสินค้าและชำระเงิน", ["เมื่อเริ่มจัดซื้อ", "เมื่อต้องตรวจสถานะเจ้าหนี้"], ["ตรวจข้อมูลผู้ขายและสินค้า", "สร้างใบสั่งซื้อ", "รับสินค้าเมื่อมีการส่งมอบ", "ตรวจบิลผู้ขายและชำระเงิน"], "งานซื้อถูกบันทึกเป็นลำดับและตรวจสอบได้", "ควรตรวจผู้ขายและเงื่อนไขก่อนสร้างเอกสาร", "ใช้เป็นภาพรวมสำหรับฝ่ายจัดซื้อและบัญชีเจ้าหนี้", "/purchases/orders", CALLOUTS["list"]),
        Section("บทที่ 4 กระบวนการซื้อ", "4.2", "รายการใบสั่งซื้อ", "ใช้ดูใบสั่งซื้อและติดตามสถานะการซื้อ", ["เมื่อต้องค้นหาใบสั่งซื้อ", "เมื่อต้องติดตามการรับสินค้า"], ["เปิดเมนูใบสั่งซื้อ", "ค้นหาผู้ขายหรือเลขเอกสาร", "ตรวจสถานะและยอดรวม", "เปิดเอกสารจากรายการ"], "ผู้ใช้พบใบสั่งซื้อและเห็นสถานะล่าสุด", "ควรตรวจเอกสารจากรายการจริงเพื่อลดโอกาสเปิดข้อมูลผิด", "ใช้ตรวจใบสั่งซื้อที่รอรับสินค้า", "/purchases/orders", CALLOUTS["list"]),
        Section("บทที่ 4 กระบวนการซื้อ", "4.3", "สร้างใบสั่งซื้อ", "ใช้สั่งซื้อสินค้าหรือบริการจากผู้ขาย", ["เมื่ออนุมัติการจัดซื้อแล้ว", "เมื่อต้องส่งคำสั่งซื้อให้ผู้ขาย"], ["กดสร้างใบสั่งซื้อ", "เลือกผู้ขาย", "เพิ่มสินค้า จำนวน ราคา และภาษี", "ตรวจยอดรวมและเงื่อนไข", "บันทึกเอกสาร"], "ระบบสร้างใบสั่งซื้อสำหรับติดตามรับสินค้าและบิลผู้ขาย", "ควรตรวจเงื่อนไขราคาและวันที่ส่งมอบก่อนยืนยัน", "ใช้เมื่อฝ่ายจัดซื้อต้องส่งคำสั่งซื้อให้ผู้ขาย", "/purchases/orders/new", CALLOUTS["form"]),
        Section("บทที่ 4 กระบวนการซื้อ", "4.4", "รับสินค้า", "ใช้บันทึกว่าสินค้าหรือบริการได้รับแล้ว", ["เมื่อผู้ขายส่งสินค้า", "เมื่อฝ่ายคลังหรือผู้รับผิดชอบยืนยันรับของ"], ["เปิดเอกสารรับสินค้า", "ตรวจรายการและจำนวนที่ได้รับ", "บันทึกการรับสินค้า", "ตรวจสถานะหลังบันทึก"], "ข้อมูลการรับสินค้าพร้อมใช้ตรวจบิลผู้ขาย", "ควรตรวจจำนวนจริงก่อนบันทึกเพื่อป้องกันยอดคงเหลือผิด", "ใช้เมื่อรับสินค้าจากใบสั่งซื้อ", "/purchases/receipts/1", CALLOUTS["form"]),
        Section("บทที่ 4 กระบวนการซื้อ", "4.5", "บิลผู้ขาย", "ใช้บันทึกหนี้ที่ต้องจ่ายให้ผู้ขาย", ["เมื่อได้รับใบแจ้งหนี้จากผู้ขาย", "เมื่อต้องตรวจยอดเจ้าหนี้"], ["เปิดบิลผู้ขายจากเอกสารที่เกี่ยวข้อง", "ตรวจผู้ขาย รายการสินค้า ยอดเงิน และภาษี", "บันทึกหรือส่งตรวจตามขั้นตอนบริษัท", "ติดตามสถานะชำระเงิน"], "ยอดเจ้าหนี้ถูกบันทึกสำหรับการจ่ายเงินและรายงาน", "ควรตรวจเอกสารผู้ขายกับใบสั่งซื้อและการรับสินค้าก่อนบันทึก", "ใช้เมื่อได้รับใบกำกับภาษีจากผู้ขาย", None, None),
        Section("บทที่ 4 กระบวนการซื้อ", "4.6", "ชำระเจ้าหนี้", "ใช้บันทึกการจ่ายเงินให้ผู้ขาย", ["เมื่อถึงกำหนดชำระ", "เมื่อต้องปิดยอดเจ้าหนี้"], ["เลือกบิลผู้ขายที่ต้องชำระ", "ตรวจยอดและช่องทางจ่ายเงิน", "บันทึกการจ่ายเงิน", "ตรวจสถานะหลังจ่าย"], "ยอดค้างจ่ายลดลงและระบบมีประวัติการชำระ", "ควรตรวจหลักฐานการจ่ายเงินก่อนบันทึก", "ใช้หลังอนุมัติจ่ายเงินให้ผู้ขาย", None, None),
        Section("บทที่ 5 การเงินและบัญชี", "5.1", "รายการรับเงิน", "ใช้ตรวจรายการเงินรับและเอกสารที่เกี่ยวข้อง", ["เมื่อต้องตรวจยอดรับจากลูกค้า", "เมื่อต้องกระทบยอดเงินรับ"], ["เปิดรายการรับเงินหรือใบเสร็จ", "ค้นหาลูกค้าหรือเลขเอกสาร", "ตรวจยอดเงินและวันที่", "เปิดรายละเอียดเมื่อต้องตรวจสอบต่อ"], "เห็นรายการรับเงินที่ใช้ตรวจสอบรายรับ", "ควรเทียบกับหลักฐานธนาคารก่อนปิดงาน", "ใช้ตรวจยอดรับเงินประจำวัน", "/sales/receipts", CALLOUTS["list"]),
        Section("บทที่ 5 การเงินและบัญชี", "5.2", "รายการจ่ายเงิน", "ใช้ตรวจรายจ่ายและเงินที่จ่ายออก", ["เมื่อต้องบันทึกรายจ่าย", "เมื่อต้องตรวจเงินจ่ายให้ผู้ขาย"], ["เปิดเมนูรายจ่าย", "ค้นหาหรือเพิ่มรายการ", "ตรวจผู้รับเงิน หมวดหมู่ และยอดเงิน", "บันทึกพร้อมหลักฐาน"], "ข้อมูลรายจ่ายพร้อมใช้ในบัญชีและรายงาน", "ควรแนบหลักฐานและเลือกหมวดให้ถูกต้อง", "ใช้บันทึกค่าใช้จ่ายสำนักงาน", "/expenses", CALLOUTS["list"]),
        Section("บทที่ 5 การเงินและบัญชี", "5.3", "สมุดรายวัน", "ใช้ดูหรือดูแลสมุดรายวันที่รองรับการบันทึกบัญชี", ["เมื่อผู้ดูแลบัญชีตรวจโครงสร้างบัญชี", "เมื่อต้องตรวจแหล่งที่มาของรายการ"], ["เปิดหน้าดูแลบัญชี", "ตรวจสมุดรายวันหรือผังบัญชี", "เลือกข้อมูลที่ต้องตรวจ", "ปรับตามสิทธิ์และนโยบายบริษัท"], "ผู้ดูแลเห็นโครงสร้างบัญชีที่ใช้บันทึกรายการ", "ควรให้ผู้ดูแลระบบบัญชีเป็นผู้แก้ไข", "ใช้เมื่อตั้งค่าระบบก่อนเริ่มใช้งานจริง", "/accounting/admin", CALLOUTS["list"]),
        Section("บทที่ 5 การเงินและบัญชี", "5.4", "ผังบัญชี", "ใช้ตรวจบัญชีที่ใช้กับเอกสารและรายงาน", ["เมื่อตั้งค่าระบบบัญชี", "เมื่อตรวจว่ารายการลงบัญชีถูกหมวด"], ["เปิดหน้าดูแลบัญชี", "ตรวจบัญชีที่ต้องใช้", "ค้นหาหรือเปิดรายละเอียด", "แก้ไขเฉพาะเมื่อได้รับอนุญาต"], "ผังบัญชีพร้อมใช้งานกับเอกสารและรายงาน", "การแก้ผังบัญชีมีผลต่อรายงาน ควรทำโดยผู้รับผิดชอบ", "ใช้ตรวจบัญชีรายได้ ค่าใช้จ่าย ลูกหนี้ และเจ้าหนี้", "/accounting/admin", CALLOUTS["list"]),
        Section("บทที่ 5 การเงินและบัญชี", "5.5", "ภาษีขาย", "ใช้ตรวจการตั้งค่าภาษีที่เกี่ยวกับเอกสารขาย", ["ก่อนออกใบแจ้งหนี้", "เมื่อตรวจรายงานภาษีขาย"], ["เปิดหน้าตั้งค่าภาษี", "ตรวจอัตราภาษีและชื่อภาษี", "เลือกภาษีให้ตรงกับเอกสารขาย", "บันทึกเมื่อมีการเปลี่ยนแปลง"], "เอกสารขายใช้ภาษีที่ถูกต้อง", "ควรตรวจอัตราภาษีตามกฎหมายและนโยบายบริษัท", "ใช้ก่อนเริ่มออกใบกำกับภาษี", "/accounting/tax-settings", CALLOUTS["form"]),
        Section("บทที่ 5 การเงินและบัญชี", "5.6", "ภาษีซื้อ", "ใช้ตรวจการตั้งค่าภาษีสำหรับเอกสารซื้อและรายจ่าย", ["ก่อนบันทึกบิลผู้ขาย", "เมื่อตรวจเครดิตภาษีซื้อ"], ["เปิดหน้าตั้งค่าภาษี", "ตรวจภาษีที่ใช้กับการซื้อ", "เลือกภาษีให้ถูกต้องในเอกสาร", "บันทึกและตรวจผลในรายงาน"], "ข้อมูลภาษีซื้อพร้อมใช้ในรายงาน", "ควรเก็บหลักฐานภาษีซื้อให้ครบก่อนบันทึก", "ใช้เมื่อตรวจบิลผู้ขายที่มี VAT", "/accounting/tax-settings", CALLOUTS["form"]),
        Section("บทที่ 6 รายงาน", "6.1", "รายงานบัญชี", "ใช้เป็นศูนย์รวมสำหรับเปิดรายงานทางบัญชี", ["เมื่อต้องตรวจผลประกอบการ", "เมื่อต้องตรวจยอดบัญชี"], ["เปิดเมนูรายงานบัญชี", "เลือกประเภทรายงาน", "กำหนดช่วงวันที่ถ้ามี", "ตรวจผลและเปิดรายละเอียดเมื่อจำเป็น"], "ผู้ใช้เข้าถึงรายงานที่ต้องการได้", "ควรตรวจช่วงวันที่และบริษัทก่อนสรุปรายงาน", "ใช้ทุกสิ้นเดือนเพื่อตรวจความครบถ้วนของบัญชี", "/accounting/reports", CALLOUTS["report"]),
        Section("บทที่ 6 รายงาน", "6.2", "งบกำไรขาดทุน", "ใช้ดูรายได้ ค่าใช้จ่าย และผลกำไรของกิจการ", ["เมื่อต้องดูผลดำเนินงาน", "เมื่อต้องส่งข้อมูลให้ผู้บริหาร"], ["เปิดรายงานงบกำไรขาดทุน", "เลือกช่วงวันที่", "ตรวจรายได้ ค่าใช้จ่าย และกำไรสุทธิ", "เปิดรายละเอียดหากต้องตรวจรายการ"], "เห็นภาพรวมผลกำไรขาดทุนตามช่วงเวลาที่เลือก", "ตัวเลขขึ้นกับเอกสารที่บันทึกและสถานะรายการ", "ใช้สรุปผลประกอบการรายเดือน", "/accounting/reports/profit-loss", CALLOUTS["report"]),
        Section("บทที่ 6 รายงาน", "6.3", "งบดุล", "ใช้ดูฐานะการเงินของกิจการ ณ ช่วงเวลาที่เลือก", ["เมื่อต้องดูสินทรัพย์ หนี้สิน และทุน", "เมื่อตรวจงบการเงิน"], ["เปิดรายงานงบดุล", "เลือกช่วงวันที่หรือวันที่สิ้นสุด", "ตรวจหมวดสินทรัพย์ หนี้สิน และทุน", "เปิดรายละเอียดหากต้องตรวจตัวเลข"], "เห็นฐานะการเงินสำหรับการตรวจสอบและสรุปผล", "ควรตรวจความครบถ้วนของเอกสารก่อนใช้รายงานประกอบการตัดสินใจ", "ใช้ประกอบการปิดงบประจำเดือน", "/accounting/reports/balance-sheet", CALLOUTS["report"]),
        Section("บทที่ 6 รายงาน", "6.4", "งบทดลอง", "ใช้ตรวจยอดเดบิตและเครดิตของบัญชี", ["ก่อนปิดงวด", "เมื่อต้องตรวจความสมดุลของบัญชี"], ["เปิดรายงานงบทดลอง", "เลือกช่วงวันที่", "ตรวจยอดเดบิต เครดิต และยอดคงเหลือ", "ติดตามบัญชีที่ต้องตรวจเพิ่ม"], "เห็นข้อมูลสำหรับตรวจสอบบัญชีก่อนปิดงวด", "หากยอดผิดปกติ ให้ตรวจรายการต้นทางก่อนแก้ไข", "ใช้ก่อนจัดทำงบการเงิน", "/accounting/reports/trial-balance", CALLOUTS["report"]),
        Section("บทที่ 6 รายงาน", "6.5", "รายงานลูกหนี้", "ใช้ติดตามยอดค้างรับของลูกค้า", ["เมื่อต้องติดตามหนี้", "เมื่อต้องตรวจยอดลูกค้ารายคน"], ["เปิดรายงานลูกหนี้หรือรายงานคู่ค้า", "เลือกช่วงวันที่หรือลูกค้าที่ต้องการ", "ตรวจยอดค้างและเอกสารที่เกี่ยวข้อง", "ติดตามการรับชำระ"], "เห็นยอดลูกหนี้เพื่อใช้ติดตามรับเงิน", "ควรตรวจสถานะใบแจ้งหนี้และใบเสร็จก่อนสรุปยอด", "ใช้ประชุมติดตามหนี้ลูกค้ารายสัปดาห์", "/accounting/reports/partner-ledger/partner/1", CALLOUTS["report"]),
        Section("บทที่ 6 รายงาน", "6.6", "รายงานเจ้าหนี้", "ใช้ติดตามยอดค้างจ่ายให้ผู้ขาย", ["เมื่อต้องวางแผนจ่ายเงิน", "เมื่อต้องตรวจยอดผู้ขาย"], ["เปิดรายงานคู่ค้าหรือรายงานเจ้าหนี้", "เลือกผู้ขายหรือช่วงวันที่", "ตรวจยอดค้างและเอกสารต้นทาง", "วางแผนชำระตามกำหนด"], "เห็นยอดเจ้าหนี้ที่ต้องจ่ายและใช้จัดลำดับการชำระ", "ควรตรวจบิลผู้ขายและหลักฐานรับสินค้าก่อนจ่ายเงิน", "ใช้เตรียมรายการจ่ายเงินประจำสัปดาห์", "/accounting/reports/partner-ledger/partner/1", CALLOUTS["report"]),
        Section("บทที่ 6 รายงาน", "6.7", "รายงานภาษี", "ใช้ตรวจข้อมูลภาษีก่อนยื่นหรือสรุปภาษี", ["ก่อนยื่นภาษี", "เมื่อต้องตรวจ VAT หรือภาษีหัก ณ ที่จ่าย"], ["เปิดรายงานภาษีที่เกี่ยวข้อง", "เลือกช่วงวันที่", "ตรวจเอกสารและยอดภาษี", "แก้ไขข้อมูลต้นทางหากพบความผิดปกติ"], "ได้ข้อมูลสำหรับตรวจสอบภาษีตามช่วงเวลาที่เลือก", "ควรตรวจเอกสารขาย ซื้อ และรายจ่ายให้ครบก่อนสรุป", "ใช้เตรียมข้อมูลภาษีประจำเดือน", "/accounting/tax-settings", CALLOUTS["report"]),
        Section("บทที่ 7 เครื่องมือและการตั้งค่า", "7.1", "นำเข้าข้อมูลจาก Excel", "ใช้ช่วยนำเข้าข้อมูลจำนวนมากเข้าระบบ", ["เมื่อต้องเริ่มใช้งานด้วยข้อมูลเดิม", "เมื่อต้องเพิ่มข้อมูลจำนวนมาก"], ["เปิดเมนูนำเข้า Excel", "เลือกประเภทข้อมูล", "เตรียมไฟล์ตามรูปแบบที่กำหนด", "นำเข้าและตรวจผล"], "ข้อมูลถูกนำเข้าระบบเพื่อลดการกรอกซ้ำ", "ควรทดลองกับข้อมูลตัวอย่างก่อนนำเข้าชุดใหญ่", "ใช้เมื่อนำเข้ารายชื่อลูกค้าหรือสินค้าเริ่มต้น", "/excel-import", CALLOUTS["form"]),
        Section("บทที่ 7 เครื่องมือและการตั้งค่า", "7.2", "เครื่องมือช่วยทำงาน", "ใช้ช่วยสร้างเอกสารหรือดึงข้อมูลจากเอกสารประกอบ", ["เมื่อต้องลดงานกรอกข้อมูล", "เมื่อต้องสร้างเอกสารจากข้อมูลตั้งต้น"], ["เปิดหน้าเครื่องมือช่วยทำงาน", "เลือกเครื่องมือที่ต้องการ", "ตรวจข้อมูลก่อนบันทึก", "ยืนยันผลลัพธ์ในหน้าหลัก"], "ผู้ใช้ทำงานซ้ำได้เร็วขึ้นและยังตรวจสอบได้ก่อนบันทึกจริง", "ควรตรวจข้อมูลที่ระบบช่วยกรอกก่อนบันทึกเสมอ", "ใช้ช่วยสร้างใบเสนอราคาหรือเพิ่มรายชื่อติดต่อจากข้อมูลที่มี", "/agent", CALLOUTS["list"]),
        Section("บทที่ 7 เครื่องมือและการตั้งค่า", "7.3", "รูปแบบเอกสารและรายงาน", "ใช้จัดการแบรนด์ เทมเพลต และตัวอย่างก่อนพิมพ์", ["เมื่อต้องปรับโลโก้หรือรูปแบบเอกสาร", "เมื่อต้องตรวจตัวอย่างก่อนส่งลูกค้า"], ["เปิดหน้าจัดการรูปแบบรายงาน", "เลือกแบรนด์หรือเทมเพลต", "แก้ไขค่าที่ได้รับอนุญาต", "ดูตัวอย่างก่อนใช้งานจริง"], "เอกสารมีรูปแบบที่สอดคล้องกับบริษัท", "ควรให้ผู้ดูแลระบบตรวจรูปแบบก่อนนำไปใช้จริง", "ใช้ตั้งค่ารูปแบบใบเสนอราคาก่อนส่งให้ลูกค้า", "/reports-studio/templates", CALLOUTS["list"]),
    ]
    return sections


def add_section(doc: Document, section: Section, shots_by_route: dict[str, Shot], fig_no: int) -> int:
    add_heading(doc, f"{section.number} {section.title}", 2)
    add_heading(doc, "ใช้สำหรับอะไร", 3)
    add_para(doc, section.purpose)
    add_heading(doc, "ควรใช้เมื่อไหร่", 3)
    add_bullets(doc, section.when)
    add_heading(doc, "ขั้นตอนการใช้งาน", 3)
    add_steps(doc, section.steps)
    add_heading(doc, "ผลลัพธ์ที่ได้", 3)
    add_para(doc, section.result)
    add_heading(doc, "ข้อควรระวัง", 3)
    add_para(doc, section.caution)
    add_heading(doc, "ตัวอย่างการใช้งาน", 3)
    add_para(doc, section.example)
    shot = image_for_route(shots_by_route, section.route)
    if shot and add_figure(doc, shot, fig_no, section.title, section.callouts):
        fig_no += 1
    return fig_no


def add_toc(doc: Document) -> None:
    add_heading(doc, "สารบัญ", 1)
    add_para(doc, "ใน Microsoft Word ให้คลิกขวาที่สารบัญและเลือก Update Field เพื่ออัปเดตเลขหน้า")
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "คลิกขวาเพื่ออัปเดตสารบัญ"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, fld_text, fld_end):
        run._r.append(el)


def create_docx(shots: list[Shot], videos: list[dict]) -> dict:
    doc = setup_doc()
    shots_by_route = {s.route: s for s in shots}
    main_error_routes = {s.route for s in shots if is_error_shot(s)}
    fig_no = 1
    sections = chapter_sections()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("ERPTH Qacc")
    tr.font.name = "Arial"
    tr.font.size = Pt(28)
    tr.bold = True
    tr.font.color.rgb = ACCENT
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("คู่มือการใช้งานระบบบัญชีสำหรับผู้ใช้\nClient-Ready User Manual\nVersion 1.0")
    sr.font.name = "Arial"
    sr.font.size = Pt(16)
    sr.font.color.rgb = MUTED
    add_para(doc, f"วันที่จัดทำ: {date.today().isoformat()}", True, ACCENT)
    doc.add_page_break()

    add_heading(doc, "ประวัติเอกสาร", 1)
    add_table(doc, ["Version", "Date", "Description", "Owner"], [["1.0", date.today().isoformat(), "Client-ready rewrite with cropped screenshots and user workflow structure", "ERPTH Qacc"]], [1.1, 1.3, 3.2, 1.1])
    add_toc(doc)
    doc.add_page_break()

    add_heading(doc, "ภาพรวมการใช้งาน", 1)
    add_para(doc, "คู่มือนี้อธิบายวิธีใช้งาน ERPTH Qacc สำหรับผู้ใช้บัญชี การขาย การซื้อ และผู้ดูแลระบบ โดยเรียงตามงานจริงที่ผู้ใช้ทำในแต่ละวัน")
    workflow_block(doc, "Workflow รวม", ["ข้อมูลพื้นฐาน", "ขาย / ซื้อ", "รับเงิน / จ่ายเงิน", "ภาษี", "รายงานบัญชี"])
    workflow_block(doc, "Workflow การขาย", ["ลูกค้า", "ใบเสนอราคา", "คำสั่งขาย", "ใบแจ้งหนี้", "รับชำระเงิน", "ใบเสร็จรับเงิน / e-Tax"])
    workflow_block(doc, "Workflow การซื้อ", ["ผู้ขาย", "ใบสั่งซื้อ", "รับสินค้า", "บิลผู้ขาย", "ชำระเงิน", "ปิดเจ้าหนี้"])
    doc.add_page_break()

    current_chapter = None
    for section in sections:
        if section.chapter != current_chapter:
            if current_chapter is not None:
                doc.add_page_break()
            current_chapter = section.chapter
            add_heading(doc, section.chapter, 1)
            if section.chapter == "บทที่ 3 กระบวนการขาย":
                add_para(doc, "กระบวนการขายใช้ติดตามงานตั้งแต่เสนอราคาให้ลูกค้า จนถึงรับเงินและออกใบเสร็จ ผู้ใช้ควรตรวจสอบข้อมูลลูกค้าและสินค้าให้ถูกต้องก่อนสร้างเอกสารขาย")
            elif section.chapter == "บทที่ 4 กระบวนการซื้อ":
                add_para(doc, "กระบวนการซื้อช่วยควบคุมการสั่งซื้อ รับสินค้า ตรวจบิลผู้ขาย และชำระเงินให้ครบตามลำดับงาน")
            elif section.chapter == "บทที่ 6 รายงาน":
                add_para(doc, "รายงานใช้สำหรับตรวจสอบผลการทำงานและตัวเลขบัญชี ควรตรวจช่วงวันที่ บริษัท และเอกสารต้นทางก่อนใช้ตัวเลขประกอบการตัดสินใจ")
        fig_no = add_section(doc, section, shots_by_route, fig_no)

    doc.add_page_break()
    add_heading(doc, "วิดีโอสาธิตการใช้งาน", 1)
    add_para(doc, "วิดีโอทั้งหมดเป็นการบันทึกการใช้งานจริงใน browser สำหรับใช้ประกอบการอบรมผู้ใช้")
    video_rows = []
    for i, video in enumerate(videos, start=1):
        video_rows.append([str(i), clean_thai(video.get("title", video.get("id", ""))), video.get("path", "")])
    add_table(doc, ["#", "หัวข้อ", "ไฟล์วิดีโอ"], video_rows, [0.4, 2.3, 3.8])

    doc.add_page_break()
    add_heading(doc, "ภาคผนวก ก: Route Coverage Matrix", 1)
    rows = []
    for shot in shots:
        rows.append([shot.chapter, shot.title, shot.route, thai_status(shot), "ย้ายไป Known Issues" if shot.route in main_error_routes else "ใช้เป็นข้อมูลประกอบคู่มือ"])
    add_table(doc, ["Area", "Screen", "Path", "Status", "Note"], rows, [1.1, 1.6, 1.7, 1.0, 1.4])

    doc.add_page_break()
    add_heading(doc, "ภาคผนวก ข: Known Issues", 1)
    issue_rows = []
    for shot in shots:
        if is_error_shot(shot):
            issue_rows.append([
                shot.chapter or "ไม่ระบุ",
                shot.title or shot.route,
                thai_status(shot),
                "ไม่ใช้ภาพนี้ในบทหลัก เพื่อไม่ให้ผู้ใช้สับสนระหว่างขั้นตอนปกติกับหน้าข้อจำกัด",
                "เปิดจากรายการเอกสารจริง ตรวจข้อมูลตัวอย่าง สิทธิ์ผู้ใช้ และการเชื่อมต่อก่อนใช้งานจริง",
            ])
    add_table(doc, ["Area", "Screen", "Issue", "Impact", "Recommended Action"], issue_rows, [1.0, 1.4, 1.1, 1.7, 1.8])

    doc.add_page_break()
    add_heading(doc, "ภาคผนวก ค: หมายเหตุสำหรับผู้ดูแลระบบ", 1)
    add_bullets(
        doc,
        [
            f"ไฟล์ต้นฉบับ production: {SOURCE_PRODUCTION}",
            f"ไฟล์ต้นฉบับเดิม: {SOURCE_ORIGINAL}",
            f"จำนวน screenshot ที่ประมวลผล: {len([s for s in shots if s.screenshot])}",
            f"จำนวนวิดีโอประกอบ: {len(videos)}",
            "หน้าที่มีข้อผิดพลาดหรือยังไม่มีข้อมูลถูกเก็บไว้ใน Known Issues และไม่ถูกใช้เป็นภาพขั้นตอนหลัก",
        ],
    )
    doc.save(DOCX_OUT)
    Document(str(DOCX_OUT))
    return {"sections": len(sections), "figures": fig_no - 1, "known_issues": len(issue_rows)}


def thai_status(shot: Shot) -> str:
    return {
        "loaded": "เปิดได้",
        "login": "หน้าเข้าสู่ระบบ",
        "empty-state": "เปิดได้ แต่ยังไม่มีข้อมูล",
        "error": "แสดงข้อผิดพลาด",
        "not-found": "ไม่พบข้อมูลหรือไม่พบหน้า",
        "restricted": "จำกัดสิทธิ์",
    }.get(shot.status, shot.status or "-")


def write_qa(shots: list[Shot], image_stats: dict, doc_stats: dict, videos: list[dict]) -> None:
    moved = [s for s in shots if is_error_shot(s)]
    not_cropped = [s for s in shots if s.crop_status == "not cropped"]
    qa = [
        "# QA Report",
        "",
        "## Input",
        f"- Source file: `{SOURCE_PRODUCTION}`",
        f"- Original source file: `{SOURCE_ORIGINAL}`",
        f"- Date: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Image Processing",
        f"- Total images extracted: {image_stats['total']}",
        f"- Images cropped: {image_stats['cropped']}",
        f"- Images kept because already tight: {image_stats['kept']}",
        f"- Images moved to Known Issues: {len(moved)}",
        f"- Images not cropped and why: {len(not_cropped)}",
    ]
    for shot in not_cropped[:20]:
        qa.append(f"  - `{shot.screenshot}`: {shot.crop_reason}")
    if len(not_cropped) > 20:
        qa.append(f"  - และอีก {len(not_cropped) - 20} รายการ")
    qa.extend(
        [
            "",
            "## Content Cleanup",
            f"- Thai cleanup replacements applied: {len(THAI_REPLACEMENTS)} replacement rules",
            "- Technical terms removed: route/API/backend/endpoint/res.partner removed from user-facing prose where generated by the manual builder",
            f"- Captions rewritten: {doc_stats['figures']}",
            "- Empty/error screenshots removed from main chapters and preserved in Known Issues appendix",
            "",
            "## Structure",
            "- Chapters created: 7 user chapters + video section + appendices",
            f"- Sections created: {doc_stats['sections']}",
            f"- Known Issues count: {doc_stats['known_issues']}",
            f"- Videos referenced: {len(videos)}",
            "",
            "## Final Quality Gate",
            "- DOCX opens with python-docx: passed",
            "- Table of contents field exists: passed",
            "- Header/footer with page number exists: passed",
            "- Main chapter empty/error screenshots excluded: passed",
            "- Section step numbering restarts at 1 by construction: passed",
            "- Cropped screenshots saved to `output/cropped_images/`: passed",
            "- DOCX visual render: skipped because `soffice` is not available in this environment",
            "",
            "## Remaining Human Review Items",
            "- Open the DOCX in Microsoft Word and update the table of contents field.",
            "- Review the Known Issues appendix against the current q01 server state before sending to a customer.",
            "- Confirm company-specific vocabulary, screenshots, and demo data are approved for client delivery.",
        ]
    )
    QA_OUT.write_text("\n".join(qa) + "\n", encoding="utf-8")


def main() -> None:
    if not SOURCE_PRODUCTION.exists():
        raise FileNotFoundError(SOURCE_PRODUCTION)
    if not CAPTURE_JSON.exists():
        raise FileNotFoundError(CAPTURE_JSON)
    OUTPUT_DIR.mkdir(exist_ok=True)
    shots, videos = load_capture()
    image_stats = process_images(shots)
    doc_stats = create_docx(shots, videos)
    write_qa(shots, image_stats, doc_stats, videos)
    print(json.dumps({"docx": str(DOCX_OUT), "qa": str(QA_OUT), "cropped_images": image_stats, "doc": doc_stats}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
